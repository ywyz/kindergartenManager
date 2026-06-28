"""课程审议 Word 导出器。"""

from __future__ import annotations

import io
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.core.logging import get_logger

logger = get_logger(__name__)

TEMPLATE_PATH = (
    Path(__file__).resolve().parents[3] / "templates" / "coursereviewactivity.docx"
)


def _get_value(record: Any, key: str) -> str:
    if isinstance(record, dict):
        value = record.get(key, "")
    else:
        value = getattr(record, key, "")
    return "" if value is None else str(value)


def _get_bool(record: Any, key: str) -> bool:
    if isinstance(record, dict):
        value = record.get(key, False)
    else:
        value = getattr(record, key, False)
    return bool(value)


def _set_font(run, size_pt: float = 11, bold: bool = False) -> None:
    run.font.name = "宋体"
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    r_fonts = rpr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), "宋体")


def _clear_cell(cell) -> None:
    for para in list(cell.paragraphs):
        para._element.getparent().remove(para._element)
    if not cell.paragraphs:
        cell.add_paragraph()


def _write_cell(cell, text: str) -> None:
    _clear_cell(cell)
    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    lines = (text or "").splitlines() or [""]
    _set_font(para.add_run(lines[0]))
    for line in lines[1:]:
        extra = cell.add_paragraph()
        _set_font(extra.add_run(line))


def _clear_paragraph(paragraph: Paragraph) -> None:
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


def _write_paragraph(paragraph: Paragraph, text: str) -> None:
    _clear_paragraph(paragraph)
    lines = (text or "").splitlines() or [""]
    for idx, line in enumerate(lines):
        if idx:
            paragraph.add_run().add_break()
        _set_font(paragraph.add_run(line), size_pt=12)


def _iter_body_blocks(doc: Document):
    for child in doc._element.body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, doc)
        elif child.tag.endswith("}tbl"):
            yield Table(child, doc)


def _trim_sample_content(doc: Document) -> None:
    """保留空白表格与原稿/二次修改占位，删除模板后方示例内容。"""
    children = list(doc._element.body.iterchildren())
    revision_idx: int | None = None
    for idx, child in enumerate(children):
        if child.tag.endswith("}p"):
            paragraph = Paragraph(child, doc)
            if paragraph.text.strip() == "二次修改":
                revision_idx = idx
                break

    if revision_idx is None:
        return

    keep_until = revision_idx
    if revision_idx + 1 < len(children) and children[revision_idx + 1].tag.endswith("}p"):
        keep_until = revision_idx + 1

    body = doc._element.body
    for child in children[keep_until + 1 :]:
        body.remove(child)


def _paragraph_after(doc: Document, marker: str) -> Paragraph | None:
    paragraphs = doc.paragraphs
    for idx, paragraph in enumerate(paragraphs):
        if paragraph.text.strip() == marker:
            if idx + 1 < len(paragraphs):
                return paragraphs[idx + 1]
            return doc.add_paragraph()
    return None


def _fill_template(doc: Document, record: Any) -> None:
    _trim_sample_content(doc)
    if not doc.tables:
        return

    table = doc.tables[0]
    rows = table.rows
    if len(rows) < 8:
        return

    _write_cell(rows[0].cells[1], _get_value(record, "activity_name"))
    _write_cell(rows[1].cells[3], _get_value(record, "grade"))
    _write_cell(rows[1].cells[5], _get_value(record, "child_count"))
    _write_cell(rows[2].cells[1], _get_value(record, "teacher_name"))
    _write_cell(rows[2].cells[3], _get_value(record, "teacher_name"))
    _write_cell(rows[2].cells[5], _get_value(record, "activity_time"))

    _write_cell(rows[4].cells[1], _get_value(record, "activity_goal"))
    _write_cell(rows[5].cells[1], _get_value(record, "activity_prep"))
    _write_cell(rows[6].cells[1], _get_value(record, "activity_process"))

    goal_adjusted = _get_bool(record, "goal_adjusted")
    goal_adjustment = _get_value(record, "goal_adjustment") if goal_adjusted else ""
    _write_cell(
        rows[4].cells[4],
        "\n".join(
            [
                "（结合实际情况对目标的调整）",
                f"保持原设计不变（{'√' if not goal_adjusted else ''}）",
                f"有所调整（{'√' if goal_adjusted else ''}）",
                "调整内容：" + (f"\n{goal_adjustment}" if goal_adjustment else ""),
            ]
        ),
    )

    prep_adjusted = _get_bool(record, "prep_adjusted")
    prep_adjustment = _get_value(record, "prep_adjustment") if prep_adjusted else ""
    _write_cell(
        rows[5].cells[4],
        "\n".join(
            [
                f"保持原设计，不变（{'√' if not prep_adjusted else ''}）",
                f"有所调整（{'√' if prep_adjusted else ''}）",
                "调整内容：" + (f"\n{prep_adjustment}" if prep_adjustment else ""),
            ]
        ),
    )

    _write_cell(
        rows[6].cells[4],
        "调整内容：\n" + _get_value(record, "process_adjustment"),
    )
    _write_cell(rows[7].cells[1], _get_value(record, "review_reason"))

    original_paragraph = _paragraph_after(doc, "附教案：原稿")
    if original_paragraph is not None:
        _write_paragraph(original_paragraph, _get_value(record, "lesson_plan_original"))

    revised_paragraph = _paragraph_after(doc, "二次修改")
    if revised_paragraph is not None:
        _write_paragraph(revised_paragraph, _get_value(record, "revised_lesson_plan"))


def _build_from_scratch(record: Any) -> Document:
    doc = Document()
    title = doc.add_paragraph("课程审议记录表")
    if title.runs:
        _set_font(title.runs[0], size_pt=16, bold=True)

    table = doc.add_table(rows=8, cols=6)
    labels = [
        ("活动名称", _get_value(record, "activity_name"), "活动形式", "集体√、区域、亲子、小组、其他", "", ""),
        ("实施单位", "南通市崇川区樾府幼儿园", "班级（年龄段）", _get_value(record, "grade"), "幼儿人数", _get_value(record, "child_count")),
        ("组织教师", _get_value(record, "teacher_name"), "记录人", _get_value(record, "teacher_name"), "活动时间", _get_value(record, "activity_time")),
        ("原稿", "", "", "", "修订", ""),
        ("活动目标", _get_value(record, "activity_goal"), "", "", "", ""),
        ("活动准备", _get_value(record, "activity_prep"), "", "", "", ""),
        ("活动过程记录", _get_value(record, "activity_process"), "", "", "", ""),
        ("简要说明课程审议后调整的理由", _get_value(record, "review_reason"), "", "", "", ""),
    ]
    for row_idx, row_data in enumerate(labels):
        for col_idx, value in enumerate(row_data):
            _write_cell(table.rows[row_idx].cells[col_idx], value)

    _write_cell(
        table.rows[4].cells[4],
        "有所调整（√）\n调整内容：\n" + _get_value(record, "goal_adjustment")
        if _get_bool(record, "goal_adjusted")
        else "保持原设计不变（√）\n有所调整（）\n调整内容：",
    )
    _write_cell(
        table.rows[5].cells[4],
        "有所调整（√）\n调整内容：\n" + _get_value(record, "prep_adjustment")
        if _get_bool(record, "prep_adjusted")
        else "保持原设计，不变（√）\n有所调整（）\n调整内容：",
    )
    _write_cell(table.rows[6].cells[4], "调整内容：\n" + _get_value(record, "process_adjustment"))

    doc.add_paragraph("附教案：原稿")
    original = doc.add_paragraph()
    _write_paragraph(original, _get_value(record, "lesson_plan_original"))
    doc.add_paragraph("二次修改")
    revised = doc.add_paragraph()
    _write_paragraph(revised, _get_value(record, "revised_lesson_plan"))
    return doc


def export_course_review_activity(
    record: Any,
    template_path: Path | None = None,
) -> bytes:
    """导出课程审议 Word 文档。"""
    tpl = template_path or TEMPLATE_PATH
    if tpl.exists():
        doc = Document(str(tpl))
        _fill_template(doc, record)
    else:
        logger.warning("课程审议模板缺失，降级从零构建", extra={"template": str(tpl)})
        doc = _build_from_scratch(record)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
