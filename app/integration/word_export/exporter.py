"""Word 文档导出服务。

主方案：打开模板 `templates/teacherplan.docx`，按其既有单元格结构填充内容，
将同一项目的子字段分别写入对应单元格（而非全部塞进一格）。

模板表格结构（19 行 2 列，左列为标题，右列为内容；多行通过纵向合并归属同一标题）：

  R0   第N周                       （整行合并）
  R1   月 日 周X                   （整行合并）
  R2   晨间活动 | 体能大循环 / 集体游戏 / 自主游戏
  R3            | 重点指导 / 活动目标 / 指导要点
  R4   晨间谈话 | 话题
  R5            | 问题设计
  R6   集体活动 | 活动主题
  R7            | 活动目标
  R8            | 活动准备
  R9            | 活动重点
  R10           | 活动难点
  R11           | 活动过程（差异标红）
  R12  室内区域 | 游戏区域
  R13           | 重点指导 / 活动目标 / 指导要点
  R14           | 支持策略
  R15  户外游戏 | 游戏区域
  R16           | 重点观察 / 活动目标 / 指导要点
  R17           | 支持策略
  R18  一日活动反思 | （内容）

若模板文件缺失，降级为从零构建一张简化表格（_export_from_scratch）。
"""
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from app.core.logging import get_logger
from app.core.models.daily_plan import DailyPlan

logger = get_logger(__name__)

_RED = RGBColor(0xFF, 0x00, 0x00)

# 模板路径：app/integration/word_export/exporter.py → 上溯 3 层到项目根
TEMPLATE_PATH = Path(__file__).resolve().parents[3] / "templates" / "teacherplan.docx"

# AI 生成文本中可识别的子字段标签
_KNOWN_LABELS = (
    "体能大循环",
    "集体游戏",
    "自主游戏",
    "重点指导",
    "重点观察",
    "活动目标",
    "指导要点",
    "游戏区域",
    "支持策略",
    "谈话主题",
    "话题",
    "问题设计",
)


def _set_font(
    run,
    size_pt: float = 11,
    bold: bool = False,
    color: RGBColor | None = None,
) -> None:
    """统一设置 run 字体（宋体）、字号、粗体与颜色。

    必须显式设置东亚字体属性，否则中文可能出现乱码。
    """
    run.font.name = "宋体"
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), "宋体")


def _parse_fields(text: str | None) -> dict[str, str]:
    """将 AI 生成的行结构文本解析为 {标签: 内容} 字典。

    识别形如 "标签：内容" 的行；其后不含标签的行（如编号 1./2./3.）
    归属上一个标签的内容（以换行拼接）。无法识别的内容被忽略。

    Returns:
        有序 dict（按出现顺序），键为标签，值为内容（可含换行）。
    """
    result: dict[str, str] = {}
    current: str | None = None
    for raw in (text or "").splitlines():
        line = raw.strip()
        if not line:
            continue
        matched: str | None = None
        rest = ""
        for label in _KNOWN_LABELS:
            if line.startswith(label + "：") or line.startswith(label + ":"):
                matched = label
                rest = line[len(label) + 1:].strip()
                break
        if matched:
            current = matched
            result[current] = rest
        elif current is not None:
            result[current] = (result[current] + "\n" + line).strip("\n")
    return result


def _reset_cell(cell) -> None:
    """清空单元格内所有段落（保留 tcPr 等属性）。"""
    for para in list(cell.paragraphs):
        para._element.getparent().remove(para._element)


def _fill_fields_cell(cell, items: list[tuple[str, str]]) -> None:
    """以 "标签：内容" 形式填充单元格，多个标签各占独立段落。

    Args:
        cell: 目标单元格。
        items: [(显示标签, 内容)]，内容可含换行（编号项各占一行）。
    """
    _reset_cell(cell)
    for label, value in items:
        value = (value or "").strip()
        lines = value.split("\n") if value else [""]
        para = cell.add_paragraph()
        _set_font(para.add_run(f"{label}："), bold=True)
        if lines[0]:
            _set_font(para.add_run(lines[0]))
        for extra in lines[1:]:
            extra_para = cell.add_paragraph()
            _set_font(extra_para.add_run(extra))
    if not cell.paragraphs:
        cell.add_paragraph()


def _fill_plain_cell(cell, text: str, size_pt: float = 11, bold: bool = False) -> None:
    """清空单元格并写入纯文本（按换行分段）。"""
    _reset_cell(cell)
    lines = (text or "").split("\n") or [""]
    for line in lines:
        para = cell.add_paragraph()
        _set_font(para.add_run(line), size_pt=size_pt, bold=bold)
    if not cell.paragraphs:
        cell.add_paragraph()


def _fill_process_cell(
    cell,
    diff_result: list[dict],
    adapted: str | None,
) -> None:
    """填充集体活动「活动过程」单元格，按差异结果标红。"""
    _reset_cell(cell)
    # "活动过程：" 标签独占一段，后续内容各行另起段落
    label_para = cell.add_paragraph()
    _set_font(label_para.add_run("活动过程："), bold=True)
    if diff_result:
        for item in diff_result:
            color = _RED if item.get("changed") else None
            lines = item.get("text", "").split("\n")
            para = cell.add_paragraph()
            _set_font(para.add_run(lines[0]), color=color)
            for extra in lines[1:]:
                if not extra.strip():
                    continue
                para = cell.add_paragraph()
                _set_font(para.add_run(extra), color=color)
    elif adapted:
        lines = adapted.split("\n")
        para = cell.add_paragraph()
        _set_font(para.add_run(lines[0]))
        for extra in lines[1:]:
            if not extra.strip():
                continue
            para = cell.add_paragraph()
            _set_font(para.add_run(extra))


def _fill_template(doc: Document, daily_plan: DailyPlan, diff_result: list[dict]) -> None:
    """按模板既有单元格结构填充各字段。"""
    table = doc.tables[0]
    rows = table.rows

    def rcell(idx: int):
        """右列内容单元格。"""
        return rows[idx].cells[1]

    # ── R0：第N周（整行合并，写左列即整行）
    week_text = (
        f"第 {daily_plan.week_number} 周" if daily_plan.week_number else "第 — 周"
    )
    _fill_plain_cell(rows[0].cells[0], week_text, size_pt=14, bold=True)

    # ── R1：月 日 周X（整行合并）
    d = daily_plan.plan_date
    date_text = f"{d.month} 月 {d.day} 日  {daily_plan.weekday_cn}" if d else ""
    _fill_plain_cell(rows[1].cells[0], date_text, size_pt=12, bold=True)

    # ── R2/R3：晨间活动
    me = _parse_fields(daily_plan.morning_activity)
    if me:
        _fill_fields_cell(
            rcell(2),
            [
                ("体能大循环", me.get("体能大循环", "")),
                ("集体游戏", me.get("集体游戏", "")),
                ("自主游戏", me.get("自主游戏", "")),
            ],
        )
        _fill_fields_cell(
            rcell(3),
            [
                ("重点指导", me.get("重点指导", "")),
                ("活动目标", me.get("活动目标", "")),
                ("指导要点", me.get("指导要点", "")),
            ],
        )
    elif daily_plan.morning_activity:
        _fill_plain_cell(rcell(2), daily_plan.morning_activity)

    # ── R4/R5：晨间谈话
    talk = _parse_fields(daily_plan.morning_talk_topic)
    topic = talk.get("谈话主题") or talk.get("话题")
    questions = talk.get("问题设计")
    if topic is None and questions is None:
        # 旧数据/纯文本：topic 字段存整段，questions 字段单独存
        topic = daily_plan.morning_talk_topic or ""
        questions = daily_plan.morning_talk_questions or ""
    _fill_fields_cell(rcell(4), [("话题", topic or "")])
    _fill_fields_cell(rcell(5), [("问题设计", questions or "")])

    # ── R6~R11：集体活动
    _fill_fields_cell(rcell(7), [("活动目标", daily_plan.activity_goal or "")])
    _fill_fields_cell(rcell(8), [("活动准备", daily_plan.activity_prep or "")])
    _fill_fields_cell(rcell(9), [("活动重点", daily_plan.activity_key or "")])
    _fill_fields_cell(rcell(10), [("活动难点", daily_plan.activity_difficult or "")])
    _fill_process_cell(rcell(11), diff_result, daily_plan.activity_process_adapted)

    # ── R12~R14：室内区域游戏
    area = _parse_fields(daily_plan.indoor_area)
    if area:
        _fill_fields_cell(rcell(12), [("游戏区域", area.get("游戏区域", ""))])
        _fill_fields_cell(
            rcell(13),
            [
                ("重点指导", area.get("重点指导", "")),
                ("活动目标", area.get("活动目标", "")),
                ("指导要点", area.get("指导要点", "")),
            ],
        )
        if area.get("支持策略"):
            _fill_fields_cell(rcell(14), [("支持策略", area.get("支持策略", ""))])
    elif daily_plan.indoor_area:
        _fill_plain_cell(rcell(12), daily_plan.indoor_area)

    # ── R15~R17：户外游戏
    outdoor = _parse_fields(daily_plan.outdoor_activity)
    if outdoor:
        _fill_fields_cell(rcell(15), [("游戏区域", outdoor.get("游戏区域", ""))])
        _fill_fields_cell(
            rcell(16),
            [
                # AI 输出标签为「重点指导」，模板户外栏标题为「重点观察」
                ("重点观察", outdoor.get("重点观察") or outdoor.get("重点指导", "")),
                ("活动目标", outdoor.get("活动目标", "")),
                ("指导要点", outdoor.get("指导要点", "")),
            ],
        )
        if outdoor.get("支持策略"):
            _fill_fields_cell(rcell(17), [("支持策略", outdoor.get("支持策略", ""))])
    elif daily_plan.outdoor_activity:
        _fill_plain_cell(rcell(15), daily_plan.outdoor_activity)

    # ── R18：一日活动反思
    if daily_plan.daily_reflection:
        _fill_plain_cell(rows[18].cells[1], daily_plan.daily_reflection)


def export_daily_plan(daily_plan: DailyPlan, diff_result: list[dict]) -> bytes:
    """生成每日活动计划 Word 文档，返回文档字节流。

    优先使用模板 `templates/teacherplan.docx` 填充其既有单元格；模板缺失或
    填充异常时降级为从零构建简化表格。

    Args:
        daily_plan: 数据库中查询到的 DailyPlan 对象。
        diff_result: 由 diff_service.compute_diff() 计算的差异列表，
                     格式 [{"text": str, "changed": bool}]。
                     changed=True 的句子在导出文档中以红色字体显示。

    Returns:
        Word 文档的 bytes 内容（不写磁盘，由调用方决定存储方式）。
    """
    if TEMPLATE_PATH.exists():
        try:
            doc = Document(str(TEMPLATE_PATH))
            _fill_template(doc, daily_plan, diff_result)
            buf = BytesIO()
            doc.save(buf)
            return buf.getvalue()
        except Exception as exc:  # noqa: BLE001 — 模板异常时降级，保证导出可用
            logger.warning(
                "模板填充失败，降级为从零构建",
                extra={"error": f"{type(exc).__name__}: {exc}"},
            )

    return _export_from_scratch(daily_plan, diff_result)


# ──────────────────────────────────────────────────────────────────────────────
# 降级方案：从零构建简化表格（模板缺失或填充异常时使用）
# ──────────────────────────────────────────────────────────────────────────────
def _write_cell(cell, text: str, size_pt: float = 11, bold: bool = False) -> None:
    """清空单元格，写入纯文本并应用字体样式。"""
    cell.text = text or ""
    for para in cell.paragraphs:
        for run in para.runs:
            _set_font(run, size_pt=size_pt, bold=bold)


def _merge_row(table, row_idx: int) -> None:
    """横向合并指定行的两列为一格。"""
    table.rows[row_idx].cells[0].merge(table.rows[row_idx].cells[1])


def _build_collective_cell(
    cell,
    daily_plan: DailyPlan,
    diff_result: list[dict],
) -> None:
    """填充集体活动单元格，活动过程按差异结果标红。"""
    cell.text = ""
    used_first = False

    def _para() -> Any:
        nonlocal used_first
        if not used_first:
            used_first = True
            return cell.paragraphs[0]
        return cell.add_paragraph()

    for label, value in [
        ("活动目标", daily_plan.activity_goal),
        ("活动准备", daily_plan.activity_prep),
        ("活动重点", daily_plan.activity_key),
        ("活动难点", daily_plan.activity_difficult),
    ]:
        if value:
            lines = value.split("\n")
            p = _para()
            _set_font(p.add_run(f"{label}："), bold=True)
            _set_font(p.add_run(lines[0]))
            for extra in lines[1:]:
                p = _para()
                _set_font(p.add_run(extra))

    has_adapted = bool(daily_plan.activity_process_adapted)
    if diff_result or has_adapted:
        p = _para()
        _set_font(p.add_run("活动过程："), bold=True)
        if diff_result:
            for item in diff_result:
                color = _RED if item.get("changed") else None
                lines = item["text"].split("\n")
                _set_font(p.add_run(lines[0]), color=color)
                for extra in lines[1:]:
                    p = _para()
                    _set_font(p.add_run(extra), color=color)
        elif has_adapted:
            lines = (daily_plan.activity_process_adapted or "").split("\n")
            _set_font(p.add_run(lines[0]))
            for extra in lines[1:]:
                p = _para()
                _set_font(p.add_run(extra))


def _export_from_scratch(daily_plan: DailyPlan, diff_result: list[dict]) -> bytes:
    """从零构建简化单表格文档（8 行 2 列），作为模板缺失时的兜底方案。"""
    doc = Document()
    for para in list(doc.paragraphs):
        para._element.getparent().remove(para._element)

    table = doc.add_table(rows=8, cols=2)
    table.style = "Table Grid"

    _merge_row(table, 0)
    week_text = (
        f"第 {daily_plan.week_number} 周" if daily_plan.week_number else "第 — 周"
    )
    _write_cell(table.rows[0].cells[0], week_text, size_pt=14, bold=True)

    _merge_row(table, 1)
    d = daily_plan.plan_date
    date_text = f"{d.month} 月 {d.day} 日  {daily_plan.weekday_cn}" if d else ""
    _write_cell(table.rows[1].cells[0], date_text, size_pt=12, bold=True)

    _write_cell(table.rows[2].cells[0], "晨间活动", bold=True)
    _write_cell(table.rows[2].cells[1], daily_plan.morning_activity or "")

    _write_cell(table.rows[3].cells[0], "晨间谈话", bold=True)
    talk = "\n".join(
        t
        for t in [daily_plan.morning_talk_topic, daily_plan.morning_talk_questions]
        if t
    )
    _write_cell(table.rows[3].cells[1], talk)

    _write_cell(table.rows[4].cells[0], "集体活动", bold=True)
    _build_collective_cell(table.rows[4].cells[1], daily_plan, diff_result)

    _write_cell(table.rows[5].cells[0], "室内区域活动", bold=True)
    _write_cell(table.rows[5].cells[1], daily_plan.indoor_area or "")

    _write_cell(table.rows[6].cells[0], "户外游戏活动", bold=True)
    _write_cell(table.rows[6].cells[1], daily_plan.outdoor_activity or "")

    _write_cell(table.rows[7].cells[0], "一日活动反思", bold=True)
    _write_cell(table.rows[7].cells[1], daily_plan.daily_reflection or "")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ──────────────────────────────────────────────────────────────────────────────
# 批量导出：将多天计划合并为单个 Word 文档
# ──────────────────────────────────────────────────────────────────────────────

def export_batch_daily_plans(
    plans_with_diffs: list[tuple[DailyPlan, list[dict]]],
) -> bytes:
    """将多条每日计划合并导出为单个 Word 文档，返回文档字节流。

    各计划按 plan_date 升序排列，相邻计划表格之间插入一个空行段落。
    空列表时返回一个仅含空白段落的合法 docx bytes。

    使用 `copy.deepcopy` 将每个独立生成的模板表格 XML 元素追加到首个文档的
    body 中，保留模板样式（字体、边框等）。

    Args:
        plans_with_diffs: [(DailyPlan, diff_result)] 列表，顺序不限。

    Returns:
        合并后的 Word 文档 bytes。
    """
    from copy import deepcopy

    if not plans_with_diffs:
        # 空列表：返回仅含一个空白段落的合法文档
        empty_doc = Document()
        buf = BytesIO()
        empty_doc.save(buf)
        return buf.getvalue()

    # 按日期升序排列
    sorted_items = sorted(plans_with_diffs, key=lambda t: t[0].plan_date)

    use_template = TEMPLATE_PATH.exists()

    def _gen_doc(plan: DailyPlan, diff: list[dict]) -> Document:
        """生成单日计划文档对象（不转 bytes）。"""
        if use_template:
            try:
                doc = Document(str(TEMPLATE_PATH))
                _fill_template(doc, plan, diff)
                return doc
            except Exception as exc:  # noqa: BLE001
                logger.warning(
                    "批量导出：模板填充失败，降级从零构建",
                    extra={"plan_date": str(plan.plan_date), "error": f"{type(exc).__name__}: {exc}"},
                )
        # 降级：_export_from_scratch 返回 bytes，重新解析为 Document
        scratch_bytes = _export_from_scratch(plan, diff)
        return Document(BytesIO(scratch_bytes))

    # 第一个计划作为主文档
    first_plan, first_diff = sorted_items[0]
    combined_doc = _gen_doc(first_plan, first_diff)

    for plan, diff in sorted_items[1:]:
        # 追加空行段落作为间隔
        combined_doc.add_paragraph()

        # 生成当天独立文档并取第一张表格
        day_doc = _gen_doc(plan, diff)
        if not day_doc.tables:
            logger.warning(
                "批量导出：某天文档无表格，跳过",
                extra={"plan_date": str(plan.plan_date)},
            )
            continue

        # deepcopy 表格 XML 节点并追加到主文档 body
        table_elem = deepcopy(day_doc.tables[0]._element)
        combined_doc.element.body.append(table_elem)

    buf = BytesIO()
    combined_doc.save(buf)
    return buf.getvalue()
