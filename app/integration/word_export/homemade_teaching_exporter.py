"""自制教玩具 Word 导出器。"""

from __future__ import annotations

import io
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

from app.core.logging import get_logger

logger = get_logger(__name__)

TEMPLATE_PATH = Path(__file__).resolve().parents[3] / "templates" / "homemadeteaching.docx"


def _get_value(record: Any, key: str) -> str:
    if isinstance(record, dict):
        value = record.get(key, "")
    else:
        value = getattr(record, key, "")
    return "" if value is None else str(value)


def _set_font(run, size_pt: float = 11, bold: bool = False) -> None:
    run.font.name = "宋体"
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    r_fonts = rpr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), "宋体")


def _clear_cell(cell) -> None:
    for para in list(cell.paragraphs):
        para._element.getparent().remove(para._element)
    if not cell.paragraphs:
        cell.add_paragraph()


def _write_cell(cell, text: str) -> None:
    _clear_cell(cell)
    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    lines = (text or "").splitlines() or [""]
    if lines:
        _set_font(para.add_run(lines[0]))
    for line in lines[1:]:
        extra = cell.add_paragraph()
        _set_font(extra.add_run(line))


def export_homemade_teaching(
    record: Any,
    template_path: Path | None = None,
) -> bytes:
    """导出自制教玩具 Word 文档。"""
    tpl = template_path or TEMPLATE_PATH
    if tpl.exists():
        doc = Document(str(tpl))
        _fill_template(doc, record)
    else:
        logger.warning("自制教玩具模板缺失，降级从零构建", extra={"template": str(tpl)})
        doc = _build_from_scratch(record)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _fill_template(doc: Document, record: Any) -> None:
    if not doc.tables:
        return
    table = doc.tables[0]
    rows = table.rows
    mapping = [
        ("class_name", 0),
        ("teacher_name", 1),
        ("toy_name", 2),
        ("materials", 3),
        ("play_methods", 4),
    ]
    for key, row_idx in mapping:
        if row_idx < len(rows) and len(rows[row_idx].cells) > 1:
            _write_cell(rows[row_idx].cells[1], _get_value(record, key))


def _build_from_scratch(record: Any) -> Document:
    doc = Document()
    title = doc.add_paragraph("南通市崇川区樾府幼儿园教师自制教玩具情况表")
    if title.runs:
        _set_font(title.runs[0], bold=True)

    table = doc.add_table(rows=5, cols=2)
    rows_data = [
        ("班级", _get_value(record, "class_name")),
        ("姓名", _get_value(record, "teacher_name")),
        ("教玩具名称", _get_value(record, "toy_name")),
        ("所用材料", _get_value(record, "materials")),
        ("玩法", _get_value(record, "play_methods")),
    ]
    for idx, (label, value) in enumerate(rows_data):
        _write_cell(table.rows[idx].cells[0], label)
        _write_cell(table.rows[idx].cells[1], value)
    return doc
