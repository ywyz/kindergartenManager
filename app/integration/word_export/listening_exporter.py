"""一对一倾听记录 Word 导出器。

模板 `templates/OneOnOneListeningSmallSecond.docx` 含 5 个表格（健康/语言/社会/科学/艺术），
每个表格前有一个领域标题段落。单表格布局（行号以健康 31 行为例）：

  R0            元数据格（0-5列合并）：观察日期/幼儿姓名/成人数目/幼儿年龄/目标
  R1            绘画1表头：左(0-3)『幼儿绘画表征（ 月 日）』 右(4-5)『一对一倾听记录』
  R2            绘画1内容：左(0-3) 图片            右(4-5) 倾听记录（图片描述）
  R3-R4         绘画2（同结构）
  R5-R6         绘画3（同结构）
  R7            指标表头：一级指标|二级指标|评价要求|评价指标|具体标准|评价
  R8..(末-2)    指标行（每个二级指标 3 行 ★/★★/★★★）；C5『评价』列打勾 √
  R(末-1)       综合评价（C0 竖排，C1-5 合并放内容）
  R(末)         支持策略（C0 竖排，C1-5 合并放内容）

三种导出模式：
  export_combined        单幼儿 5 领域 → 1 个 docx
  export_split_by_domain 单幼儿 5 领域 → {领域: bytes}（UI 打包 zip）
  export_batch_by_domain 多幼儿按领域 → {领域: bytes}（每档含所有幼儿该领域表）
"""
from __future__ import annotations

import copy
import io
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.table import Table

from app.core.logging import get_logger

logger = get_logger(__name__)

TEMPLATE_PATH = Path(__file__).resolve().parents[3] / "templates" / "OneOnOneListeningSmallSecond.docx"

# 五大领域（模板表格顺序）
DOMAINS = ["健康", "语言", "社会", "科学", "艺术"]

# 单张绘画图片宽度（左格跨 4/6 列）
_IMG_WIDTH = Cm(6.5)


# ─── 通用单元格辅助 ──────────────────────────────────────────────────────────


def _set_font(run, size_pt: float = 11, bold: bool = False, color: RGBColor | None = None) -> None:
    """统一设置 run 宋体 + eastAsia，避免中文乱码。"""
    run.font.name = "宋体"
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    r_fonts = rpr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), "宋体")


def _clear_cell(cell) -> None:
    """清空单元格所有段落文字（保留至少 1 个空段落）。"""
    for para in cell.paragraphs:
        for elem in list(para._element):
            tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
            if tag == "r":
                para._element.remove(elem)


def _write_cell(cell, text: str | None, bold: bool = False) -> None:
    """清空并写入单元格文本，使用宋体。"""
    _clear_cell(cell)
    if not text:
        return
    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    run = para.add_run(str(text))
    _set_font(run, bold=bold)


def _put_image(cell, img_bytes: bytes | None) -> None:
    """清空单元格并在其中插入一张图片（等比按固定宽度）。"""
    if not img_bytes:
        return
    _clear_cell(cell)
    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    run = para.add_run()
    try:
        run.add_picture(io.BytesIO(img_bytes), width=_IMG_WIDTH)
    except Exception as exc:  # 非法图片字节不应中断整份导出
        logger.warning("插入图片失败，跳过", extra={"error": str(exc)})


# ─── 模板内容辅助 ────────────────────────────────────────────────────────────


def _para_text(p_element) -> str:
    """提取 <w:p> 元素的全部文本。"""
    return "".join(t.text or "" for t in p_element.iter(qn("w:t")))


def _domain_of_title(text: str) -> str | None:
    """从标题段落文本判断所属领域，非领域标题返回 None。"""
    if "领域" not in text:
        return None
    for d in DOMAINS:
        if f"{d}领域" in text:
            return d
    return None


def _norm(text: str) -> str:
    """去除全部空白字符（含换行），用于竖排标签比对。"""
    return "".join(text.split())


def _find_row_index(table: Table, predicate) -> int | None:
    """返回首个满足 predicate(row) 的行索引。"""
    for i, row in enumerate(table.rows):
        if predicate(row):
            return i
    return None


def _write_metadata(cell, record: dict, dom_data: dict) -> None:
    """在元数据格各标签行冒号后追加值（保留模板标签格式）。"""
    year = dom_data.get("obs_year") or ""
    month = dom_data.get("obs_month") or ""
    obs_date = f"{year}年{month}月" if year or month else ""
    label_values = {
        "观察日期": obs_date,
        "幼儿姓名": str(record.get("child_name") or ""),
        "成人数目": str(record.get("adult_count") or ""),
        "幼儿年龄": str(record.get("child_age") or ""),
        "目标": str(dom_data.get("goals") or ""),
    }
    for para in cell.paragraphs:
        ptext = para.text.replace("\u3000", "").replace(" ", "").strip()
        for label, value in label_values.items():
            if ptext.startswith(label) and value:
                run = para.add_run(value)
                _set_font(run)
                break


def _fill_date_in_header(cell, d) -> None:
    """将绘画表头『幼儿绘画表征（ 月 日）』填入具体月日（保留首个 run 格式）。"""
    if d is None:
        return
    filled = f"幼儿绘画表征（{d.month}月{d.day}日）"
    para = cell.paragraphs[0] if cell.paragraphs else None
    if para is None:
        return
    if para.runs:
        para.runs[0].text = filled
        for r in para.runs[1:]:
            r.text = ""
    else:
        run = para.add_run(filled)
        _set_font(run)


def _fill_domain_table(table: Table, record: dict, dom_data: dict) -> None:
    """将一个领域的全部数据填入对应表格。"""
    rows = table.rows
    if not rows:
        return

    # R0 元数据
    _write_metadata(rows[0].cells[0], record, dom_data)

    # 指标表头行
    header_idx = _find_row_index(table, lambda r: r.cells[0].text.strip() == "一级指标")
    if header_idx is None:
        header_idx = min(7, len(rows))

    # 绘画/倾听记录块：R1/R3/R5 表头 + R2/R4/R6 内容
    dates = [dom_data.get("date_1"), dom_data.get("date_2"), dom_data.get("date_3")]
    images = dom_data.get("images") or []  # list[(bytes, description)]
    slot = 0
    hdr_idx = 1
    while hdr_idx + 1 < header_idx and slot < 3:
        content_idx = hdr_idx + 1
        _fill_date_in_header(rows[hdr_idx].cells[0], dates[slot])
        if slot < len(images):
            img_bytes, desc = images[slot]
            _put_image(rows[content_idx].cells[0], img_bytes)
            if desc:
                _write_cell(rows[content_idx].cells[4], desc)
        slot += 1
        hdr_idx += 2

    # 综合评价 / 支持策略 行
    zh_idx = _find_row_index(table, lambda r: _norm(r.cells[0].text) == "综合评价")
    st_idx = _find_row_index(table, lambda r: _norm(r.cells[0].text) == "支持策略")
    end_indicator = zh_idx if zh_idx is not None else len(rows)

    # 指标打勾：第 i 个二级指标 3 行（★/★★/★★★），标记达成星级行的 C5
    indicators = sorted(
        dom_data.get("indicators") or [],
        key=lambda x: x.get("sort_order", 0),
    )
    for i, ind in enumerate(indicators):
        stars = ind.get("stars", 3)
        try:
            stars = max(1, min(3, int(stars)))
        except (TypeError, ValueError):
            stars = 3
        target = header_idx + 1 + i * 3 + (stars - 1)
        if target < end_indicator:
            _write_cell(rows[target].cells[5], "√")

    if zh_idx is not None:
        _write_cell(rows[zh_idx].cells[1], dom_data.get("evaluation"))
    if st_idx is not None:
        _write_cell(rows[st_idx].cells[1], dom_data.get("support_strategy"))


def _find_domain(domains: list[dict], domain: str) -> dict | None:
    """从领域 payload 列表中取指定领域。"""
    for d in domains:
        if d.get("domain") == domain:
            return d
    return None


def _map_tables_to_domains(doc) -> list[tuple[str | None, Table]]:
    """按文档顺序将每个表格映射到其前置标题段落判定的领域。"""
    result: list[tuple[str | None, Table]] = []
    current_domain: str | None = None
    table_idx = 0
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            d = _domain_of_title(_para_text(child))
            if d:
                current_domain = d
        elif child.tag == qn("w:tbl"):
            if table_idx < len(doc.tables):
                result.append((current_domain, doc.tables[table_idx]))
            table_idx += 1
    return result


def _save_bytes(doc) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─── 块抽取 / 移除（用于拆分与批量）─────────────────────────────────────────


def _extract_block(doc, domain: str):
    """从文档中抽取某领域的 (标题段落元素, 表格元素)，未找到返回 (None, None)。"""
    children = list(doc.element.body.iterchildren())
    for i, child in enumerate(children):
        if child.tag == qn("w:p") and _domain_of_title(_para_text(child)) == domain:
            for nxt in children[i + 1:]:
                if nxt.tag == qn("w:tbl"):
                    return child, nxt
                if nxt.tag == qn("w:p") and _domain_of_title(_para_text(nxt)):
                    break
    return None, None


def _remove_all_blocks(doc) -> None:
    """移除 body 内全部段落与表格（保留 sectPr）。"""
    body = doc.element.body
    for child in list(body.iterchildren()):
        tag = child.tag
        if tag in (qn("w:p"), qn("w:tbl")):
            body.remove(child)


def _build_domain_doc(domain: str, children: list[tuple[dict, list[dict]]], tpl: Path) -> bytes:
    """构建仅含指定领域的文档，按 children 顺序为每个幼儿追加一份该领域表格。"""
    if not tpl.exists():
        logger.warning("一对一倾听模板缺失，降级从零构建", extra={"template": str(tpl)})
        return _build_from_scratch(domain, children)

    pristine = Document(str(tpl))
    title_el, tbl_el = _extract_block(pristine, domain)
    if tbl_el is None:
        logger.warning("模板中未找到领域块，降级", extra={"domain": domain})
        return _build_from_scratch(domain, children)

    doc = Document(str(tpl))
    _remove_all_blocks(doc)
    body = doc.element.body
    sectPr = body.find(qn("w:sectPr"))

    for record, domains in children:
        dom_data = _find_domain(domains, domain)
        if dom_data is None:
            continue
        new_title = copy.deepcopy(title_el)
        new_tbl = copy.deepcopy(tbl_el)
        if sectPr is not None:
            sectPr.addprevious(new_title)
            sectPr.addprevious(new_tbl)
        else:
            body.append(new_title)
            body.append(new_tbl)
        _fill_domain_table(doc.tables[-1], record, dom_data)

    return _save_bytes(doc)


# ─── 公开导出接口 ────────────────────────────────────────────────────────────


def export_combined(record: dict, domains: list[dict], template_path: Path | None = None) -> bytes:
    """单幼儿 5 领域合并导出为 1 个 docx。"""
    tpl = template_path or TEMPLATE_PATH
    if not tpl.exists():
        logger.warning("一对一倾听模板缺失，降级从零构建", extra={"template": str(tpl)})
        doc = Document()
        for d in DOMAINS:
            dom_data = _find_domain(domains, d)
            if dom_data is not None:
                _scratch_domain(doc, d, record, dom_data)
        return _save_bytes(doc)

    doc = Document(str(tpl))
    for domain, table in _map_tables_to_domains(doc):
        if domain is None:
            continue
        dom_data = _find_domain(domains, domain)
        if dom_data is not None:
            _fill_domain_table(table, record, dom_data)
    return _save_bytes(doc)


def export_split_by_domain(
    record: dict, domains: list[dict], template_path: Path | None = None
) -> dict[str, bytes]:
    """单幼儿按领域拆分为多个 docx，返回 {领域: bytes}（仅含有数据的领域）。"""
    tpl = template_path or TEMPLATE_PATH
    result: dict[str, bytes] = {}
    for d in DOMAINS:
        if _find_domain(domains, d) is not None:
            result[d] = _build_domain_doc(d, [(record, domains)], tpl)
    return result


def export_batch_by_domain(
    children: list[tuple[dict, list[dict]]],
    domain_order: list[str] | None = None,
    template_path: Path | None = None,
) -> dict[str, bytes]:
    """多幼儿按领域批量导出，返回 {领域: bytes}（每档含所有幼儿该领域表）。

    Args:
        children: [(record, domains), ...] 每个幼儿的数据。
        domain_order: 领域顺序（默认模板顺序）。
        template_path: 模板路径。
    """
    tpl = template_path or TEMPLATE_PATH
    order = domain_order or DOMAINS
    result: dict[str, bytes] = {}
    for d in order:
        kids = [(rec, doms) for rec, doms in children if _find_domain(doms, d) is not None]
        if kids:
            result[d] = _build_domain_doc(d, kids, tpl)
    return result


# ─── 模板缺失降级 ────────────────────────────────────────────────────────────


def _scratch_domain(doc, domain: str, record: dict, dom_data: dict) -> None:
    """模板缺失时，为一个领域构建简化内容（标题 + 基本字段）。"""
    title = doc.add_paragraph(f'"一对一倾听"观察记录 — {domain}领域')
    _set_font(title.runs[0] if title.runs else title.add_run(), bold=True)
    lines = [
        f"幼儿姓名：{record.get('child_name', '')}  年龄：{record.get('child_age', '')}",
        f"观察日期：{dom_data.get('obs_year', '')}年{dom_data.get('obs_month', '')}月",
        f"目标：{dom_data.get('goals', '') or ''}",
        f"综合评价：{dom_data.get('evaluation', '') or ''}",
        f"支持策略：{dom_data.get('support_strategy', '') or ''}",
    ]
    for line in lines:
        p = doc.add_paragraph()
        _set_font(p.add_run(line))


def _build_from_scratch(domain: str, children: list[tuple[dict, list[dict]]]) -> bytes:
    """模板缺失时构建单领域多幼儿简化文档。"""
    doc = Document()
    for record, domains in children:
        dom_data = _find_domain(domains, domain)
        if dom_data is not None:
            _scratch_domain(doc, domain, record, dom_data)
            doc.add_paragraph("")
    return _save_bytes(doc)
