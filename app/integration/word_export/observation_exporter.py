"""游戏观察记录 Word 导出器。

主方案：打开模板 `templates/ObservationRecord.docx`，
替换标题中的 'xx' 为大环境，将各字段写入对应单元格。

模板表格结构（8 行 × 8 列）：

  R0  班  级 | (班级) | 日    期 | (日期) | ... | 起止时间 | 起止时间 | (时间)
  R1  成人数目 | (成人) | 儿童数目 | (儿童) | ... | 观察者 | 观察者 | (观察者)
  R2  幼儿姓名 | (姓名×3合并) | ... | 幼儿年龄 | 幼儿年龄 | (年龄) | ...
  R3  观察环境 | (内容)
  R4  观察目标 | (内容)
  R5  观察记录 | (文字 + 图片横排)
  R6  评价分析 | (内容)
  R7  支持策略 | (内容)

字段写入规则：
  R0: 班级→C1, 日期→C3, 起止时间→C7
  R1: 成人数目→C1, 儿童数目→C3, 观察者→C7
  R2: 幼儿姓名→C1, 幼儿年龄→C6
  R3~R7: 内容→C1

模板缺失时，降级为从零构建简化表格。
"""
from __future__ import annotations

import io
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.core.logging import get_logger

logger = get_logger(__name__)

# 模板路径（项目根/templates/ObservationRecord.docx）
TEMPLATE_PATH = Path(__file__).resolve().parents[3] / "templates" / "ObservationRecord.docx"

_RED = RGBColor(0xFF, 0x00, 0x00)

# 单图最大宽度（3 张并排时约 4.5cm，1~2 张可更宽）
_IMG_WIDTHS = {1: Cm(10), 2: Cm(6), 3: Cm(4.5)}


def _set_font(run, size_pt: float = 11, bold: bool = False, color: RGBColor | None = None) -> None:
    """统一设置 run 宋体 + eastAsia，避免中文乱码。"""
    run.font.name = "宋体"
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    r_fonts = rpr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), "宋体")


def _clear_cell(cell) -> None:
    """清空单元格所有段落文字（保留至少 1 个空段落）。"""
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
        # 清空段落内直接的文本节点
        for elem in list(para._element):
            tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
            if tag == "r":
                para._element.remove(elem)


def _write_cell(cell, text: str | None, bold: bool = False) -> None:
    """清空并写入单元格文本，使用宋体。"""
    _clear_cell(cell)
    if not text:
        return
    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    run = para.add_run(text)
    _set_font(run, bold=bold)


def _add_images_to_cell(cell, images: list[bytes]) -> None:
    """在单元格现有文字之后追加一个段落横向并排插入图片。"""
    if not images:
        return

    width = _IMG_WIDTHS.get(len(images), Cm(4.5))
    para = cell.add_paragraph()
    for i, img_bytes in enumerate(images):
        run = para.add_run()
        run.add_picture(io.BytesIO(img_bytes), width=width)
        if i < len(images) - 1:
            para.add_run(" ")  # 图片间空格


def export_observation(
    observation: dict,
    images: list[bytes],
    template_path: Path | None = None,
) -> bytes:
    """将游戏观察记录导出为 Word 文档。

    Args:
        observation: 观察记录字段 dict（含所有 design.md §4.1 字段）。
        images: 压缩后图片字节列表（0~3 张，写入 R5 单元格文字下方）。
        template_path: 可选模板路径，None 时使用 TEMPLATE_PATH。

    Returns:
        Word 文档的 bytes。
    """
    tpl = template_path or TEMPLATE_PATH

    if tpl.exists():
        doc = Document(str(tpl))
        _fill_template(doc, observation, images)
    else:
        logger.warning("游戏观察模板缺失，降级从零构建", extra={"template": str(tpl)})
        doc = _build_from_scratch(observation, images)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _fill_template(doc: Document, obs: dict, images: list[bytes]) -> None:
    """填充模板文档各字段。"""
    # 1. 替换标题 paragraph 中的 'xx'
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text == "xx":
                run.text = obs.get("big_env", "户外")
                break

    # 2. 填充表格
    if not doc.tables:
        return

    table = doc.tables[0]
    rows = table.rows

    def row(n):
        return rows[n] if n < len(rows) else None

    # R0: 班级→C1, 日期→C3, 起止时间→C7
    r0 = row(0)
    if r0:
        _write_cell(r0.cells[1], str(obs.get("class_name") or ""))
        _write_cell(r0.cells[3], str(obs.get("obs_date") or ""))
        _write_cell(r0.cells[7], str(obs.get("time_range") or ""))

    # R1: 成人数目→C1, 儿童数目→C3, 观察者→C7
    r1 = row(1)
    if r1:
        _write_cell(r1.cells[1], str(obs.get("adult_count") or ""))
        _write_cell(r1.cells[3], str(obs.get("child_count") or ""))
        _write_cell(r1.cells[7], str(obs.get("observer") or ""))

    # R2: 幼儿姓名→C1, 幼儿年龄→C6（先清空示例文本）
    r2 = row(2)
    if r2:
        _write_cell(r2.cells[1], str(obs.get("child_names") or ""))
        _write_cell(r2.cells[6], str(obs.get("child_age") or ""))

    # R3~R4: 整行内容写 C1
    for row_idx, field in [(3, "game_area"), (4, "observation_goal")]:
        r = row(row_idx)
        if r:
            _write_cell(r.cells[1], str(obs.get(field) or ""))

    # R5: 观察记录文字 + 图片
    r5 = row(5)
    if r5:
        cell = r5.cells[1]
        _clear_cell(cell)
        para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        text = obs.get("observation_record") or ""
        if text:
            run = para.add_run(text)
            _set_font(run)
        _add_images_to_cell(cell, images)

    # R6: 评价分析→C1
    r6 = row(6)
    if r6:
        _write_cell(r6.cells[1], str(obs.get("evaluation_analysis") or ""))

    # R7: 支持策略→C1
    r7 = row(7)
    if r7:
        _write_cell(r7.cells[1], str(obs.get("support_strategy") or ""))


def _build_from_scratch(obs: dict, images: list[bytes]) -> Document:
    """模板缺失时，从零构建一个简化的观察记录表格（8 行 2 列）。"""
    doc = Document()
    title = doc.add_paragraph(
        f"南通市崇川区樾府幼儿园游戏观察记录（{obs.get('big_env', '户外')}）"
    )
    _set_font(title.runs[0] if title.runs else title.add_run(), bold=True)

    table = doc.add_table(rows=8, cols=2)

    rows_data = [
        ("班级/日期/起止时间", f"{obs.get('class_name','')} {obs.get('obs_date','')} {obs.get('time_range','')}"),
        ("成人/儿童/观察者", f"{obs.get('adult_count','')} / {obs.get('child_count','')} / {obs.get('observer','')}"),
        ("幼儿姓名/年龄", f"{obs.get('child_names','')} {obs.get('child_age','')}"),
        ("观察环境", obs.get("game_area") or ""),
        ("观察目标", obs.get("observation_goal") or ""),
        ("观察记录", obs.get("observation_record") or ""),
        ("评价分析", obs.get("evaluation_analysis") or ""),
        ("支持策略", obs.get("support_strategy") or ""),
    ]

    for i, (label, content) in enumerate(rows_data):
        row = table.rows[i]
        _write_cell(row.cells[0], label, bold=True)
        if i == 5:  # 观察记录行插入图片
            _write_cell(row.cells[1], content)
            _add_images_to_cell(row.cells[1], images)
        else:
            _write_cell(row.cells[1], content)

    return doc
