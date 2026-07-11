"""幼儿个体发展档案 Word 导出器。

模板 `templates/personal.docx` 含 1 个表格（8行 x 12列），布局如下：

  Row 0: 姓名(C0-2) / 性别(C3-5) / 出生年月(C6-8) / 入园时间(C9-11)
  Row 1: 身高(C0-1) / 体重(C2-4) / 胸围(C5-6) / 血色素(C7-8) / 视力(C9-11)
  Row 2: 幼儿发展情况（C0竖排标题，C1-11内容）
  Row 3: 采取措施（C0竖排标题，C1-11内容）
  Row 4: 家园联系（C0竖排标题，C1-11内容）
  Row 5: 突出表现（C0标题，C1-11内容）
  Row 6: 进步情况（C0标题，C1-11内容）
  Row 7: 保教老师寄语（C0竖排标题，C1-11内容）

导出模式：单记录导出（一个幼儿一份档案）。
"""
from __future__ import annotations

import io
from datetime import date
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.core.logging import get_logger

logger = get_logger(__name__)

TEMPLATE_PATH = Path(__file__).resolve().parents[3] / "templates" / "personal.docx"


def _set_font(run, size_pt: float = 11, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def _merge_cells(cells) -> None:
    if cells:
        first = cells[0]
        for cell in cells[1:]:
            first.merge(cell)


def _set_cell_text(cell, text: str, size_pt: float = 11, bold: bool = False) -> None:
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    _set_font(run, size_pt=size_pt, bold=bold)


def export_personal_development(data: dict) -> bytes:
    """导出幼儿个体发展档案为 Word 文档。

    Args:
        data: 档案数据字典，包含：
          child_name, gender, birth_date, enrollment_date
          height, weight, chest_circumference, hemoglobin, vision_left, vision_right
          development_status, measures_taken, home_contact
          outstanding_performance, progress, teacher_message

    Returns:
        Word 文档字节流。
    """
    doc = Document(TEMPLATE_PATH)

    if not doc.tables:
        raise ValueError("模板文件中未找到表格")

    table = doc.tables[0]

    birth_date_str = ""
    if data.get("birth_date"):
        if isinstance(data["birth_date"], date):
            birth_date_str = data["birth_date"].strftime("%Y年%m月")
        else:
            birth_date_str = str(data["birth_date"])

    enrollment_date_str = ""
    if data.get("enrollment_date"):
        if isinstance(data["enrollment_date"], date):
            enrollment_date_str = data["enrollment_date"].strftime("%Y年%m月")
        else:
            enrollment_date_str = str(data["enrollment_date"])

    _set_cell_text(table.cell(0, 0), data.get("child_name", ""), bold=True)
    _merge_cells([table.cell(0, 0), table.cell(0, 1), table.cell(0, 2)])

    _set_cell_text(table.cell(0, 3), data.get("gender", ""))
    _merge_cells([table.cell(0, 3), table.cell(0, 4), table.cell(0, 5)])

    _set_cell_text(table.cell(0, 6), birth_date_str)
    _merge_cells([table.cell(0, 6), table.cell(0, 7), table.cell(0, 8)])

    _set_cell_text(table.cell(0, 9), enrollment_date_str)
    _merge_cells([table.cell(0, 9), table.cell(0, 10), table.cell(0, 11)])

    height_str = f"{data['height']}" if data.get("height") else ""
    _set_cell_text(table.cell(1, 0), height_str)
    _merge_cells([table.cell(1, 0), table.cell(1, 1)])

    weight_str = f"{data['weight']}" if data.get("weight") else ""
    _set_cell_text(table.cell(1, 2), weight_str)
    _merge_cells([table.cell(1, 2), table.cell(1, 3), table.cell(1, 4)])

    chest_str = f"{data['chest_circumference']}" if data.get("chest_circumference") else ""
    _set_cell_text(table.cell(1, 5), chest_str)
    _merge_cells([table.cell(1, 5), table.cell(1, 6)])

    hemoglobin_str = f"{data['hemoglobin']}" if data.get("hemoglobin") else ""
    _set_cell_text(table.cell(1, 7), hemoglobin_str)
    _merge_cells([table.cell(1, 7), table.cell(1, 8)])

    vision_left = data.get("vision_left")
    vision_right = data.get("vision_right")
    vision_str = ""
    if vision_left and vision_right:
        vision_str = f"{vision_left} / {vision_right}"
    elif vision_left:
        vision_str = str(vision_left)
    elif vision_right:
        vision_str = str(vision_right)
    _set_cell_text(table.cell(1, 9), vision_str)
    _merge_cells([table.cell(1, 9), table.cell(1, 10), table.cell(1, 11)])

    content_cols = list(range(1, 12))

    _set_cell_text(table.cell(2, 1), data.get("development_status", ""))
    for col in content_cols[1:]:
        _merge_cells([table.cell(2, col)])

    _set_cell_text(table.cell(3, 1), data.get("measures_taken", ""))
    for col in content_cols[1:]:
        _merge_cells([table.cell(3, col)])

    _set_cell_text(table.cell(4, 1), data.get("home_contact", ""))
    for col in content_cols[1:]:
        _merge_cells([table.cell(4, col)])

    _set_cell_text(table.cell(5, 1), data.get("outstanding_performance", ""))
    for col in content_cols[1:]:
        _merge_cells([table.cell(5, col)])

    _set_cell_text(table.cell(6, 1), data.get("progress", ""))
    for col in content_cols[1:]:
        _merge_cells([table.cell(6, col)])

    _set_cell_text(table.cell(7, 1), data.get("teacher_message", ""))
    for col in content_cols[1:]:
        _merge_cells([table.cell(7, col)])

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()