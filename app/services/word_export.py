"""Word 导出服务 - 使用 python-docx 填充模板并生成下载文件"""
import copy
import io
import re
from datetime import date
from pathlib import Path
from typing import List

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor, Pt
from lxml import etree

from app.config import AppConfig
from app.models.daily_plan import DailyPlan


# ---------------------------------------------------------------------------
# 颜色常量
# ---------------------------------------------------------------------------
RED = RGBColor(0xFF, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)


# ---------------------------------------------------------------------------
# 辅助函数
# ---------------------------------------------------------------------------

def _clear_cell_after_label(para, label_len: int):
    """删除段落中标签之后的所有内容（保留标签本身）"""
    # 重建该段落，只保留标签
    full_text = para.text
    label_text = full_text[:label_len] if label_len <= len(full_text) else full_text

    # 清除所有 runs
    for run in para.runs:
        run.text = ""

    # 确保第一个 run 包含标签
    if para.runs:
        para.runs[0].text = label_text
    else:
        para.add_run(label_text)


def _append_content_run(para, content: str, is_red: bool = False):
    """在段落末尾追加内容 run，可选红色"""
    run = para.add_run(content)
    if is_red:
        run.font.color.rgb = RED
    else:
        # 继承模板字体颜色（黑色）
        run.font.color.rgb = BLACK


def _set_cell_text(cell, text: str, is_red: bool = False, keep_label: bool = False):
    """
    设置单元格第一个（或唯一）段落的内容。
    keep_label=True 时保留段落原有标签文字，在之后追加 text。
    """
    para = cell.paragraphs[0]
    if keep_label:
        label = para.text  # 保留标签
        for run in para.runs:
            run.text = ""
        if para.runs:
            para.runs[0].text = label
        _append_content_run(para, text, is_red)
    else:
        para.clear()
        _append_content_run(para, text, is_red)


def _fill_labeled_para(para, content: str, is_red: bool = False):
    """
    取段落中第一个冒号后面作为填充区域，清空并写入 content。
    例：段落原文 "活动主题： " → 保留 "活动主题：" + 写入 content。
    """
    orig = para.text
    colon_idx = max(orig.find("："), orig.find(":"))
    if colon_idx == -1:
        label = orig.rstrip()
    else:
        label = orig[: colon_idx + 1]  # 包含"："

    # 清空所有 runs
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = label
    else:
        para.add_run(label)

    _append_content_run(para, content, is_red)


def _fill_multiline_cell(cell, fields: list[tuple[str, str, bool]]):
    """
    向多标签单元格填充多个字段。
    fields: [(label关键字, 内容, is_red), ...]
    通过匹配段落中的标签关键字来定位。
    """
    for para in cell.paragraphs:
        text = para.text.strip()
        for label_key, content, is_red in fields:
            # 精确匹配标签关键字
            if label_key in text:
                _fill_labeled_para(para, content, is_red)
                break


# 强环节标题：中文数字"一、" / "（一）" / "第N步|环节|部分|阶段|课时"
_PROCESS_STRONG_HEADER_RE = re.compile(
    r"^\s*(?:"
    r"[一二三四五六七八九十百]+\s*[、.．:：]"
    r"|[（(]\s*[一二三四五六七八九十百]+\s*[）)]"
    r"|第[一二三四五六七八九十百\d]+\s*(?:个|步|环节|部分|阶段|课时)"
    r")"
)

# 弱环节标题：阿拉伯数字"1.""1、""(1)" —— 仅在全文不含强标题时才用作分节
_PROCESS_WEAK_HEADER_RE = re.compile(
    r"^\s*(?:"
    r"\d+\s*[、.．:：)）]"
    r"|[（(]\s*\d+\s*[）)]"
    r")"
)

# 兼容【AI修改】/[AI新增]/(AI补充) 等格式
_PROCESS_AI_TAG_RE = re.compile(r"[\[\(【]\s*AI[^\]\)】]*[\]\)】]")


def _compute_process_red_flags(lines: list[str]) -> list[bool]:
    """按"小节"判定每一行是否应红字。

    - 优先用强标题（中文数字 / "第N步" 等）划分；阿拉伯数字此时作为节内步骤，不分节。
    - 若全文无强标题，则退回用阿拉伯数字标题划分。
    - 若仍无标题，则仅对带 AI 标记的行染色。
    """
    n = len(lines)
    flags = [False] * n
    header_idx = [i for i, ln in enumerate(lines) if _PROCESS_STRONG_HEADER_RE.search(ln)]
    if not header_idx:
        header_idx = [i for i, ln in enumerate(lines) if _PROCESS_WEAK_HEADER_RE.search(ln)]

    if not header_idx:
        for i, ln in enumerate(lines):
            if _PROCESS_AI_TAG_RE.search(ln):
                flags[i] = True
        return flags

    sections: list[tuple[int, int]] = []
    if header_idx[0] > 0:
        sections.append((0, header_idx[0]))
    for k, hi in enumerate(header_idx):
        end = header_idx[k + 1] if k + 1 < len(header_idx) else n
        sections.append((hi, end))

    for start, end in sections:
        if any(_PROCESS_AI_TAG_RE.search(lines[i]) for i in range(start, end)):
            for i in range(start, end):
                flags[i] = True
    return flags


def _fill_process_cell(cell, process_text: str, use_ai_color: bool):
    """
    活动过程专用填充函数。
    - use_ai_color=False：全部黑色
    - use_ai_color=True：按环节判定，含【AI修改】/【AI新增】等标记的整节红色，其余黑色。

    实现要点：每一行写入独立段落，单 run 显式带 <w:color> 属性，避免
    Word 渲染时 run 颜色越界、或 \\n 不被识别为换行的问题。
    """
    target_para = None
    for para in cell.paragraphs:
        if "活动过程" in para.text:
            target_para = para
            break
    if target_para is None:
        return

    # 提取并保留标签（"活动过程："）
    orig = target_para.text
    colon_idx = max(orig.find("："), orig.find(":"))
    label = orig[: colon_idx + 1] if colon_idx != -1 else orig.rstrip()
    for run in target_para.runs:
        run.text = ""
    if target_para.runs:
        target_para.runs[0].text = label
    else:
        target_para.add_run(label)

    if not process_text:
        return

    lines = process_text.split("\n")
    line_red_flags = (
        _compute_process_red_flags(lines) if use_ai_color else [False] * len(lines)
    )

    # 复制原段落的 pPr，让后续行段落继承缩进/对齐等样式
    target_p_element = target_para._p
    target_pPr = target_p_element.find(qn("w:pPr"))

    prev_element = target_p_element
    for i, line in enumerate(lines):
        new_p = OxmlElement("w:p")
        if target_pPr is not None:
            new_p.append(copy.deepcopy(target_pPr))

        new_r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        color_el = OxmlElement("w:color")
        color_el.set(qn("w:val"), "FF0000" if line_red_flags[i] else "000000")
        rPr.append(color_el)
        new_r.append(rPr)

        new_t = OxmlElement("w:t")
        new_t.text = line
        new_t.set(qn("xml:space"), "preserve")
        new_r.append(new_t)
        new_p.append(new_r)

        prev_element.addnext(new_p)
        prev_element = new_p


# ---------------------------------------------------------------------------
# 主导出函数
# ---------------------------------------------------------------------------

def export_daily_plan_word(plan: DailyPlan) -> bytes:
    """
    根据 DailyPlan 填充模板，返回 Word 文件的二进制内容（bytes）。
    调用方将此内容写入文件或通过 HTTP 响应下载。
    """
    template_path = AppConfig.WORD_TEMPLATE
    if not template_path.exists():
        raise FileNotFoundError(f"Word 模板不存在：{template_path}")

    doc = Document(str(template_path))
    t = doc.tables[0]
    ma = plan.morning_activity
    mt = plan.morning_talk
    ga = plan.group_activity
    ia = plan.indoor_area
    og = plan.outdoor_game

    # 判断哪些字段是 AI 修改过的（需要红色字体）
    modified: set[str] = set(plan.ai_modified_parts.get("fields", []))

    def is_red(field_name: str) -> bool:
        return field_name in modified

    # ----------------------------------------------------------------------
    # Row 0 - 第N周（合并单元格）
    # ----------------------------------------------------------------------
    row0_cell = t.rows[0].cells[0]
    para0 = row0_cell.paragraphs[0]
    para0.clear()
    week_text = f"第 {plan.week_number} 周" if plan.week_number else ""
    _append_content_run(para0, week_text)

    # ----------------------------------------------------------------------
    # Row 1 - 日期信息（合并单元格）
    # ----------------------------------------------------------------------
    row1_cell = t.rows[1].cells[0]
    para1 = row1_cell.paragraphs[0]
    para1.clear()
    if plan.plan_date:
        d: date = plan.plan_date
        date_text = f"{d.month} 月 {d.day} 日  {plan.day_of_week}"
    else:
        date_text = ""
    _append_content_run(para1, date_text)

    # ----------------------------------------------------------------------
    # Row 2 - 晨间活动：集体活动名称 / 自选活动名称
    # ----------------------------------------------------------------------
    row2_cell = t.rows[2].cells[1]
    # 模板段落标签关键字可能为「集体游戏/集体活动」「自主游戏/自选活动」
    _fill_multiline_cell(row2_cell, [
        ("集体", ma.group_activity_name, is_red("morning_activity")),
        ("自", ma.self_selected_name, is_red("morning_activity")),
    ])

    # ----------------------------------------------------------------------
    # Row 3 - 晨间活动：重点指导/活动目标/指导要点
    # ----------------------------------------------------------------------
    row3_cell = t.rows[3].cells[1]
    _fill_multiline_cell(row3_cell, [
        ("重点指导", ma.key_guidance, is_red("morning_activity")),
        ("活动目标", ma.activity_goal, is_red("morning_activity")),
        ("指导要点", ma.guidance_points, is_red("morning_activity")),
    ])

    # ----------------------------------------------------------------------
    # Row 4 - 晨间谈话：话题
    # ----------------------------------------------------------------------
    row4_cell = t.rows[4].cells[1]
    _fill_labeled_para(row4_cell.paragraphs[0], mt.topic, is_red("morning_talk"))

    # ----------------------------------------------------------------------
    # Row 5 - 晨间谈话：问题设计
    # ----------------------------------------------------------------------
    row5_cell = t.rows[5].cells[1]
    paras = row5_cell.paragraphs
    if paras:
        p = paras[0]
        orig = p.text
        for run in p.runs:
            run.text = ""
        if p.runs:
            p.runs[0].text = orig.rstrip()
        _append_content_run(p, f"\n{mt.questions}" if mt.questions else "", is_red("morning_talk"))
    if len(paras) > 1:
        # 第二段落清空旧内容，写入问题
        pass  # 问题设计填入第一段落后的 run 即可

    # ----------------------------------------------------------------------
    # Row 6 - 集体活动：活动主题
    # ----------------------------------------------------------------------
    _fill_labeled_para(t.rows[6].cells[1].paragraphs[0], ga.theme, is_red("group_activity_theme"))

    # ----------------------------------------------------------------------
    # Row 7 - 集体活动：活动目标
    # ----------------------------------------------------------------------
    row7_cell = t.rows[7].cells[1]
    _fill_multiline_cell(row7_cell, [
        ("活动目标", ga.goal, is_red("group_activity_goal")),
    ])

    # ----------------------------------------------------------------------
    # Row 8 - 集体活动：活动准备
    # ----------------------------------------------------------------------
    row8_cell = t.rows[8].cells[1]
    _fill_multiline_cell(row8_cell, [
        ("活动准备", ga.preparation, is_red("group_activity_preparation")),
    ])

    # ----------------------------------------------------------------------
    # Row 9 - 集体活动：活动重点
    # ----------------------------------------------------------------------
    _fill_labeled_para(t.rows[9].cells[1].paragraphs[0], ga.key_point, is_red("group_activity_key_point"))

    # ----------------------------------------------------------------------
    # Row 10 - 集体活动：活动难点
    # ----------------------------------------------------------------------
    _fill_labeled_para(t.rows[10].cells[1].paragraphs[0], ga.difficulty, is_red("group_activity_difficulty"))

    # ----------------------------------------------------------------------
    # Row 11 - 集体活动：活动过程（含【AI修改】标记的行红色，其余黑色）
    # ----------------------------------------------------------------------
    row11_cell = t.rows[11].cells[1]
    _fill_process_cell(row11_cell, ga.process, is_red("group_activity_process"))

    # ----------------------------------------------------------------------
    # Row 12 - 室内区域游戏：游戏区域
    # ----------------------------------------------------------------------
    _fill_labeled_para(t.rows[12].cells[1].paragraphs[0], ia.game_area, is_red("indoor_area"))

    # ----------------------------------------------------------------------
    # Row 13 - 室内区域游戏：重点指导/活动目标/指导要点
    # ----------------------------------------------------------------------
    _fill_multiline_cell(t.rows[13].cells[1], [
        ("重点指导", ia.key_guidance, is_red("indoor_area")),
        ("活动目标", ia.activity_goal, is_red("indoor_area")),
        ("指导要点", ia.guidance_points, is_red("indoor_area")),
    ])

    # ----------------------------------------------------------------------
    # Row 14 - 室内区域游戏：支持策略
    # ----------------------------------------------------------------------
    _fill_multiline_cell(t.rows[14].cells[1], [
        ("支持策略", ia.support_strategy, is_red("indoor_area")),
    ])

    # ----------------------------------------------------------------------
    # Row 15 - 户外游戏：游戏区域
    # ----------------------------------------------------------------------
    _fill_labeled_para(t.rows[15].cells[1].paragraphs[0], og.game_area, is_red("outdoor_game"))

    # ----------------------------------------------------------------------
    # Row 16 - 户外游戏：重点观察/活动目标/指导要点
    # ----------------------------------------------------------------------
    _fill_multiline_cell(t.rows[16].cells[1], [
        ("重点观察", og.key_guidance, is_red("outdoor_game")),
        ("活动目标", og.activity_goal, is_red("outdoor_game")),
        ("指导要点", og.guidance_points, is_red("outdoor_game")),
    ])

    # ----------------------------------------------------------------------
    # Row 17 - 户外游戏：支持策略
    # ----------------------------------------------------------------------
    _fill_multiline_cell(t.rows[17].cells[1], [
        ("支持策略", og.support_strategy, is_red("outdoor_game")),
    ])

    # ----------------------------------------------------------------------
    # Row 18 - 一日活动反思
    # ----------------------------------------------------------------------
    row18_cell = t.rows[18].cells[1]
    if row18_cell.paragraphs:
        p = row18_cell.paragraphs[0]
        p.clear()
        _append_content_run(p, plan.daily_reflection)

    # ----------------------------------------------------------------------
    # 保存到内存并返回 bytes
    # ----------------------------------------------------------------------
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def save_export_to_file(plan: DailyPlan) -> Path:
    """
    将导出的 Word 写入 exports/ 目录，返回文件路径。
    文件名格式：{年级}{班级}_{YYYY-MM-DD}_{第N周周X}.docx
    """
    export_dir = AppConfig.EXPORT_DIR
    export_dir.mkdir(parents=True, exist_ok=True)

    content = export_daily_plan_word(plan)
    if plan.plan_date:
        d = plan.plan_date
        date_part = f"{d.year}-{d.month:02d}-{d.day:02d}"
        grade_class = f"{plan.grade}{plan.class_name}" if (plan.grade or plan.class_name) else ""
        week_part = ""
        if plan.week_number:
            day_of_week = plan.day_of_week or ""
            week_part = f"_第{plan.week_number}周{day_of_week}"
        if grade_class:
            filename = f"{grade_class}_{date_part}{week_part}.docx"
        else:
            filename = f"{date_part}{week_part}.docx"
    else:
        filename = "daily_plan_export.docx"

    file_path = export_dir / filename
    file_path.write_bytes(content)
    return file_path


def export_merged_plans(plans: List[DailyPlan], author_name: str = "教师") -> Path:
    """
    将多份计划合并导出为单个 Word 文件，每份计划占一页（分页符分隔）。
    plans 应已按日期升序排列。
    返回导出文件路径，文件名为 "{author_name}备课笔记.docx"。
    """
    if not plans:
        raise ValueError("没有可导出的计划")

    export_dir = AppConfig.EXPORT_DIR
    export_dir.mkdir(parents=True, exist_ok=True)

    # 第一份用正常方式生成作为基础文档
    first_bytes = export_daily_plan_word(plans[0])
    merged_doc = Document(io.BytesIO(first_bytes))

    # 后续每份：生成独立文档，将其表格复制到合并文档中（分页符分隔）
    for plan in plans[1:]:
        plan_bytes = export_daily_plan_word(plan)
        sub_doc = Document(io.BytesIO(plan_bytes))

        # 添加分页符
        merged_doc.add_page_break()

        # 复制子文档的表格到合并文档
        for table in sub_doc.tables:
            # 深拷贝表格 XML 到合并文档
            new_tbl = copy.deepcopy(table._tbl)
            merged_doc.element.body.append(new_tbl)

    filename = f"{author_name}备课笔记.docx"
    file_path = export_dir / filename
    buf = io.BytesIO()
    merged_doc.save(buf)
    file_path.write_bytes(buf.getvalue())
    return file_path
