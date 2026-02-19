"""
Word文档操作模块
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from .models import WORD_FONT_NAME, WORD_FONT_SIZE, WORD_INDENT_FIRST_LINE

# 预编译正则表达式以提高性能
_PUNCTUATION_PATTERN = re.compile(r'([。？！.?!])')


def apply_run_style(run):
    """设置运行样式为仿宋小四"""
    run.font.name = WORD_FONT_NAME
    run.font.size = Pt(WORD_FONT_SIZE)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), WORD_FONT_NAME)


def normalize_label(label):
    """标准化标签：去除空格、冒号和换行符"""
    # 先去除换行符和回车
    normalized = label.replace("\n", "").replace("\r", "")
    # 去除所有冒号（中英文）
    normalized = normalized.replace("：", "").replace(":", "")
    # 去除前后空格
    normalized = normalized.strip()
    return normalized


def normalize_multiline_text(text):
    """规范化多行文本，统一换行符并处理转义字符"""
    if text is None:
        return ""
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    if "\\n" in normalized or "\\r" in normalized:
        normalized = (
            normalized.replace("\\r\\n", "\n")
            .replace("\\n", "\n")
            .replace("\\r", "\n")
        )
    return normalized


def split_by_punctuation(text):
    """按句号、问号、叹号拆分为多行（保留中英文标点）"""
    if not text:
        return []
    # 统一换行符
    text = normalize_multiline_text(text)
    lines = []
    for line in text.splitlines():
        # 按句号、问号、叹号拆分
        parts = _PUNCTUATION_PATTERN.split(line)
        current = ""
        for i, part in enumerate(parts):
            current += part
            # 如果是标点符号且不是最后一个元素
            if part in '。？！.?!' and current.strip():
                lines.append(current.strip())
                current = ""
        if current.strip():
            lines.append(current.strip())
    return lines if lines else [""]


def set_cell_text(cell, text):
    """设置单元格文字（仅用于简单文本，如周次/日期）"""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    apply_run_style(run)


def set_cell_content(cell, text, indent=True, split_sentences=False):
    """设置单元格内容，按换行拆分为段落"""
    cell.text = ""
    if split_sentences:
        lines = split_by_punctuation(text)
    else:
        normalized = normalize_multiline_text(text)
        lines = normalized.splitlines() if normalized else [""]
    for i, line in enumerate(lines):
        if not line.strip():
            continue
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        if indent:
            p.paragraph_format.first_line_indent = Pt(WORD_INDENT_FIRST_LINE)
        run = p.add_run(line)
        apply_run_style(run)


def append_by_labels(
    cell,
    label_to_text,
    append_unmatched=True,
    context_parent=None,
):
    """
    根据标签追加内容到单元格
    
    Args:
        cell: Word表格单元格
        label_to_text: {标签名: 内容} 字典
        append_unmatched: 是否添加未匹配的新标签
        context_parent: 上下文父字段名（用于智能匹配子字段）
    """
    original_lines = cell.text.splitlines()
    if not original_lines:
        original_lines = [""]
    
    matched = set()
    
    # 清空单元格
    cell.text = ""
    
    # 重建原有内容
    for i, line in enumerate(original_lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        run = p.add_run(line)
        apply_run_style(run)
        
        # 检查该行是否匹配标签
        line_strip = line.strip()
        if "：" in line_strip:
            line_label = normalize_label(line_strip.split("：")[0])
        else:
            line_label = normalize_label(line_strip)
        
        if not line_label:
            continue
            
        # 智能查找对应的值
        extra = smart_lookup(label_to_text, line_label, context_parent)
        
        if extra and line_label not in matched:
            # 找到匹配的标签，创建新段落追加内容
            matched.add(line_label)
                
            # 按句号拆分（适用于目标、重点等）
            split_sentences = line_label in [
                "活动目标", "活动准备", "活动重点",
                "活动难点", "指导要点", "问题设计"
            ]
            if split_sentences:
                parts = split_by_punctuation(extra)
            else:
                parts = normalize_multiline_text(extra).splitlines()
            for part in parts:
                if part.strip():
                    new_p = cell.add_paragraph()
                    new_p.paragraph_format.first_line_indent = Pt(
                        WORD_INDENT_FIRST_LINE
                    )
                    run = new_p.add_run(part)
                    apply_run_style(run)
    
    # 添加未匹配的新标签（仅未带前缀的顶级字段）
    if append_unmatched:
        for label, extra in label_to_text.items():
            # 跳过带前缀的子字段
            if "-" in label:
                continue
            
            label_base = normalize_label(label)
            if label_base not in matched and extra:
                # 创建标签行
                p = cell.add_paragraph()
                run = p.add_run(f"{label_base}：")
                apply_run_style(run)
                
                # 创建内容段落（按句号拆分）
                split_sentences = label_base in [
                    "活动目标", "活动准备", "活动重点",
                    "活动难点", "指导要点", "问题设计"
                ]
                if split_sentences:
                    parts = split_by_punctuation(extra)
                else:
                    parts = normalize_multiline_text(extra).splitlines()
                for part in parts:
                    if part.strip():
                        new_p = cell.add_paragraph()
                        new_p.paragraph_format.first_line_indent = Pt(
                            WORD_INDENT_FIRST_LINE
                        )
                        run = new_p.add_run(part)
                        apply_run_style(run)


def flatten_plan_data(plan_data):
    """将嵌套的教案数据扁平化，保留层级关系"""
    flat_data = {}
    for key, value in plan_data.items():
        if isinstance(value, dict):
            # 子字段加父字段前缀，同时也保留不带前缀的（按FIELD_ORDER顺序优先）
            for sub_key, sub_value in value.items():
                if sub_value:
                    prefixed_key = f"{key}-{sub_key}"
                    flat_data[prefixed_key] = sub_value
        else:
            if value:
                flat_data[key] = value
    return flat_data


def smart_lookup(label_to_text, target_label, context_parent=None):
    """智能查找标签对应的值，支持带前缀匹配"""
    # 1. 如果有上下文父字段，优先匹配带前缀的
    if context_parent:
        prefixed = f"{context_parent}-{target_label}"
        if prefixed in label_to_text:
            return label_to_text[prefixed]
    
    # 2. 尝试直接匹配
    if target_label in label_to_text:
        return label_to_text[target_label]
    
    # 3. 查找所有匹配 "*-target_label" 模式的key
    for key, value in label_to_text.items():
        if "-" in key and key.endswith(f"-{target_label}"):
            return value
    
    return None


def fill_table_by_labels(table, label_to_text, content_col=1, label_col=0):
    """使用标签填充表格所有行的内容列"""
    context_parent = None
    for row in table.rows:
        if len(row.cells) <= content_col:
            continue
        # 尝试从标签列推断上下文父字段
        if len(row.cells) > label_col:
            label_text = normalize_label(row.cells[label_col].text)
            # 检查是否是父字段标签（更新上下文）
            # 需要找到与长标签匹配的父字段（考虑模板中可能的多行标签如"下午：\n户外游戏"）
            detected_parent = None
            for key in label_to_text.keys():
                if "-" in key:
                    parent_name = key.split("-")[0]
                    # 规范化比较，去除"："和换行符
                    parent_normalized = normalize_label(parent_name)
                    if parent_normalized == label_text:
                        detected_parent = parent_name
                        break
            
            if detected_parent:
                context_parent = detected_parent
        
        append_by_labels(
            row.cells[content_col],
            label_to_text,
            append_unmatched=False,
            context_parent=context_parent,
        )


def fill_by_row_labels(table, label_to_text, label_col=0, content_col=1):
    """按行标签填充表格（标签在label_col，内容填入content_col）"""
    context_parent = None
    
    for row in table.rows:
        if len(row.cells) <= max(label_col, content_col):
            continue
        
        label_text = normalize_label(row.cells[label_col].text)
        if not label_text:
            continue
        
        # 检查是否是父字段标签（更新上下文）
        # 需要找到与标签匹配的父字段（考虑模板中可能的多行标签如"下午：\n户外游戏"）
        detected_parent = None
        for key in label_to_text.keys():
            if "-" in key:
                parent_name = key.split("-")[0]
                # 规范化比较
                parent_normalized = normalize_label(parent_name)
                if parent_normalized == label_text:
                    detected_parent = parent_name
                    break
        
        if detected_parent:
            context_parent = detected_parent
        
        # 智能查找对应的值
        value = smart_lookup(label_to_text, label_text, context_parent)
        
        if value:
            # 对目标、重点等字段按句号拆分
            split_sentences = any(
                keyword in label_text
                for keyword in ["目标", "重点", "难点", "准备", "要点"]
            )
            set_cell_content(
                row.cells[content_col],
                value,
                indent=True,
                split_sentences=split_sentences,
            )


def fill_doc_by_labels(
    doc,
    plan_data,
    week_text=None,
    date_text=None,
    content_col=1,
    label_col=0,
    header_table_index=0,
):
    """
    根据标签填充Word文档
    
    Args:
        doc: Document对象
        plan_data: 教案数据字典
        week_text: 周次文本
        date_text: 日期文本
        content_col: 内容列索引
        label_col: 标签列索引
        header_table_index: 包含周次/日期的表格索引
    """
    flat_data = flatten_plan_data(plan_data)
    for index, table in enumerate(doc.tables):
        fill_table_by_labels(table, flat_data, content_col=content_col)

        if index == header_table_index:
            if week_text is not None and len(table.rows) > 0:
                set_cell_text(table.cell(0, content_col), week_text)
            if date_text is not None and len(table.rows) > 1:
                set_cell_text(table.cell(1, content_col), date_text)

        fill_by_row_labels(
            table,
            flat_data,
            label_col=label_col,
            content_col=content_col,
        )


def fill_teacher_plan(doc, plan_data, week_text, date_text):
    """
    填充教师教案文档
    
    Args:
        doc: Document对象
        plan_data: 教案数据
        week_text: 周次文本（如"第（1）周"）
        date_text: 日期文本（如"周（一） 2月26日"）
    """
    fill_doc_by_labels(
        doc,
        plan_data,
        week_text=plan_data.get("周次", week_text),
        date_text=plan_data.get("日期", date_text),
        content_col=1,
        label_col=0,
        header_table_index=0,
    )


def generate_plan_docx(
    template_path,
    plan_data,
    week_text,
    date_text,
    output_path,
):
    """
    生成教案Word文档
    
    Args:
        template_path: 模板文档路径
        plan_data: 教案数据
        week_text: 周次
        date_text: 日期
        output_path: 输出文件路径
        
    Returns:
        输出文件路径 (Path对象)
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    doc = Document(template_path)
    fill_teacher_plan(doc, plan_data, week_text, date_text)
    doc.save(output_path)
    
    return output_path
