"""T011 周/月候选模板资格门的稳定 RED。

候选资格是受控 seed/fixture 的内部校验 job，不是 TemplateCenter 的正式 Preview 或 CRUD。
"""

from dataclasses import FrozenInstanceError, fields, replace
from hashlib import sha256
from importlib import import_module
from inspect import signature
from pathlib import Path
from uuid import UUID
from zipfile import ZipFile

import pytest
from _support import (
    OFFICE_CLIENT_VERSIONS,
    RESERVED_TYPES,
    MemoryControlledSeedStore,
    MemoryExportPort,
    MemoryOfficeQualificationPort,
    MemoryQualificationEvidenceStore,
    docx_with_external_relationship,
    docx_with_macro_member,
    docx_with_safe_office_support_parts,
    docx_with_structure_profile_mismatch,
    docx_with_table_shapes,
    docx_with_text,
)
from docx import Document


def test_candidate_registry_pins_released_seed_hashes_and_exact_table_profiles():
    api = _api()

    assert tuple(
        (
            profile.document_type.value,
            profile.seed_handle.handle_id,
            profile.seed_handle.expected_sha256,
            profile.contract.required_anchors,
        )
        for profile in api.CANDIDATE_QUALIFICATION_PROFILES
    ) == (
        (
            "weekly_activity_plan",
            "controlled-weekplan-seed-v2",
            "157abf313206d94a90337807e490e0ea0ad8b72cf0d3eb6d7ef0ed6a6aa93f14",
            ("tables:word/document.xml:3x9x7",),
        ),
        (
            "monthly_theme_activity_plan",
            "controlled-monthplan-seed-v2",
            "de806aed3289f0a5f0019318aec63380f681dae3113383d47d03b363337b69d5",
            ("tables:word/document.xml:1x8x4",),
        ),
    )


def test_unique_validator_enforces_exact_candidate_table_count_rows_and_grid():
    api = _api()
    weekly = api.CANDIDATE_QUALIFICATION_PROFILES[0]
    monthly = api.CANDIDATE_QUALIFICATION_PROFILES[1]

    api.validate_upload(
        docx_with_table_shapes((9, 7), (9, 7), (9, 7)),
        "synthetic-weekly.docx",
        api.DOCX_MIME_TYPE,
        weekly.contract,
    )
    api.validate_upload(
        docx_with_table_shapes((8, 4)),
        "synthetic-monthly.docx",
        api.DOCX_MIME_TYPE,
        monthly.contract,
    )
    for content, contract in (
        (docx_with_table_shapes((9, 7)), weekly.contract),
        (docx_with_table_shapes((9, 7), (9, 7)), weekly.contract),
        (docx_with_table_shapes((9, 7), (9, 7), (9, 6)), weekly.contract),
        (docx_with_table_shapes((8, 4), (8, 4)), monthly.contract),
        (docx_with_table_shapes((8, 3)), monthly.contract),
    ):
        with pytest.raises(api.TemplateCenterError) as caught:
            api.validate_upload(
                content,
                "synthetic-candidate.docx",
                api.DOCX_MIME_TYPE,
                contract,
            )
        assert caught.value.code is api.TemplateErrorCode.VALIDATION_FAILED


def test_released_v2_candidate_files_are_hash_bound_structurally_sanitized_seeds():
    api = _api()
    repository_root = Path(__file__).resolve().parents[3]
    paths = {
        "weekly_activity_plan": repository_root / "templates" / "weekplan.docx",
        "monthly_theme_activity_plan": repository_root / "templates" / "monthplan.docx",
    }

    for profile in api.CANDIDATE_QUALIFICATION_PROFILES:
        path = paths[profile.document_type.value]
        content = path.read_bytes()
        assert sha256(content).hexdigest() == profile.seed_handle.expected_sha256
        receipt = api.validate_upload(
            content,
            path.name,
            api.DOCX_MIME_TYPE,
            profile.contract,
        )
        assert receipt.structural_profile_id == profile.profile_id
        assert receipt.structural_profile_version == profile.profile_version == 2
        with ZipFile(path) as package:
            assert not any(name.startswith("customXml/") for name in package.namelist())
            assert b'TargetMode="External"' not in b"".join(
                package.read(name)
                for name in package.namelist()
                if name.endswith(".rels")
            )
            core = package.read("docProps/core.xml")
            assert b"<dc:creator>" not in core
            assert b"<cp:lastModifiedBy>" not in core

    weekly = Document(paths["weekly_activity_plan"])
    assert [len(table.rows) for table in weekly.tables] == [9, 9, 9]
    assert all(len(table.columns) == 7 for table in weekly.tables)
    assert (
        tuple(weekly.paragraphs[index].text for index in (1, 5, 8))
        == ("主题名称：____  班级：____  周次：____  日期：____",) * 3
    )
    assert (
        tuple(weekly.paragraphs[index].text for index in (2, 6, 9))
        == ("教师：____  保育员：____",) * 3
    )
    for table in weekly.tables:
        assert table.rows[0].cells[1].text == ""
        assert all(
            not cell.text.strip() for row in table.rows[1:5] for cell in row.cells[2:]
        )
        assert all(not row.cells[1].text.strip() for row in table.rows[5:])

    monthly = Document(paths["monthly_theme_activity_plan"])
    assert monthly.paragraphs[1].text == (
        "班级：____  执行年月：____  带班老师：____  保育老师：____"
    )
    assert monthly.tables[0].rows[1].cells[0].text.strip() == (
        "本月主题：\n上月分析：\n本月重点："
    )
    for row_index in (3, 4, 5):
        assert not monthly.tables[0].rows[row_index].cells[1].text.strip()
        assert not monthly.tables[0].rows[row_index].cells[3].text.strip()
    assert not monthly.tables[0].rows[7].cells[0].text.strip()

    safety = repository_root / "templates" / "1530.docx"
    assert sha256(safety.read_bytes()).hexdigest() == (
        "e26b258921db61ac070b7ef124bab75316975d567b046c668ecb685c6ccba540"
    )
    with ZipFile(safety) as package:
        assert not any(name.startswith("customXml/") for name in package.namelist())
        assert b'TargetMode="External"' not in b"".join(
            package.read(name) for name in package.namelist() if name.endswith(".rels")
        )
        core = package.read("docProps/core.xml")
        assert b"<dc:creator>" not in core
        assert b"<cp:lastModifiedBy>" not in core
    safety_doc = Document(safety)
    assert safety_doc.paragraphs[0].text == "幼儿园“1530”安全教育记录"
    assert safety_doc.paragraphs[1].text == "班级：____  周次：____  时间：____"
    assert all(
        not safety_doc.tables[0].rows[row_index].cells[column_index].text.strip()
        for row_index in range(1, 7)
        for column_index in (2, 3)
    )


@pytest.mark.parametrize("include_directories", [False, True], ids=["parts", "dirs"])
def test_unique_validator_accepts_closed_safe_office_support_parts(
    include_directories,
):
    api = _api()
    contract = replace(
        api.CANDIDATE_QUALIFICATION_PROFILES[0].contract,
        allowed_parts=(
            "word/document.xml",
            "word/styles.xml",
            "word/numbering.xml",
            "word/settings.xml",
            "word/fontTable.xml",
            "word/theme/theme1.xml",
        ),
        required_anchors=("table:word/document.xml:19x2",),
    )

    api.validate_upload(
        docx_with_safe_office_support_parts(include_directories=include_directories),
        "synthetic-office.docx",
        api.DOCX_MIME_TYPE,
        contract,
    )


@pytest.mark.asyncio
async def test_registered_seed_handle_cannot_be_rebound_to_different_valid_bytes():
    api = _api()
    seeds = MemoryControlledSeedStore()
    seeds.register_controlled_seed(
        "controlled-weekplan-seed-v2", docx_with_text("rebound candidate seed")
    )
    job, effects = _job(api, seeds=seeds)

    with pytest.raises(api.TemplateCenterError) as caught:
        await job.qualify(
            document_type="weekly_activity_plan",
            seed_handle="controlled-weekplan-seed-v2",
            fixture=_fixture(api),
            profile_id="weekly_activity_plan-profile-v2",
        )

    assert caught.value.code is api.TemplateErrorCode.VALIDATION_FAILED
    assert effects["export"].render_calls == []
    assert effects["office"].calls == []
    assert effects["evidence"].items == []


def _api():
    return import_module("app.service.template_center")


@pytest.fixture(autouse=True)
def _use_closed_synthetic_profiles_for_unit_job_tests(monkeypatch):
    """Keep unit GREEN independent of the untracked released Office templates."""
    module = import_module("app.service.template_center.candidate_qualification")
    production_resolver = module.candidate_profile
    seed_bytes = MemoryControlledSeedStore().seeds

    def resolve(document_type, seed_handle, profile_id):
        profile = production_resolver(document_type, seed_handle, profile_id)
        expected = sha256(seed_bytes[profile.seed_handle.handle_id]).hexdigest()
        return replace(
            profile,
            seed_handle=replace(
                profile.seed_handle,
                expected_sha256=expected,
            ),
            contract=replace(
                profile.contract,
                required_anchors=("table:word/document.xml:19x2",),
            ),
        )

    monkeypatch.setattr(module, "candidate_profile", resolve)


def _job(api, *, seeds=None, export=None, office=None, evidence=None):
    seeds = seeds or MemoryControlledSeedStore()
    export = export or MemoryExportPort()
    office = office or MemoryOfficeQualificationPort()
    evidence = evidence or MemoryQualificationEvidenceStore()
    job = api.TemplateCandidateQualificationJob(
        controlled_seed_store=seeds,
        export_port=export,
        office_qualification_port=office,
        qualification_evidence_store=evidence,
    )
    return job, {
        "seeds": seeds,
        "export": export,
        "office": office,
        "evidence": evidence,
    }


def _fixture(api, *, provenance="synthetic"):
    return api.SyntheticQualificationFixture(
        fixture_id="weekly-monthly-fixture-v1",
        provenance=provenance,
        values=(("title", "仅用于候选资格的合成值"),),
    )


class _EqualityBypass:
    def __eq__(self, other):
        return True


class _StringSubclass(str):
    pass


@pytest.mark.asyncio
@pytest.mark.parametrize(
    ("document_type", "seed_handle", "profile_id"),
    [
        (_EqualityBypass(), _EqualityBypass(), _EqualityBypass()),
        (
            _StringSubclass("weekly_activity_plan"),
            _StringSubclass("controlled-weekplan-seed-v2"),
            _StringSubclass("weekly_activity_plan-profile-v2"),
        ),
    ],
    ids=["custom-equality", "str-subclass"],
)
async def test_candidate_preflight_rejects_equality_bypass_before_all_ports(
    document_type, seed_handle, profile_id
):
    api = _api()
    job, effects = _job(api)

    with pytest.raises(api.TemplateCenterError) as caught:
        await job.qualify(
            document_type=document_type,
            seed_handle=seed_handle,
            fixture=_fixture(api),
            profile_id=profile_id,
        )

    assert caught.value.code is api.TemplateErrorCode.INPUT_INVALID
    assert effects["seeds"].read_calls == []
    assert effects["export"].render_calls == []
    assert effects["export"].parse_calls == []
    assert effects["office"].calls == []
    assert effects["evidence"].items == []


def test_candidate_job_and_contracts_are_closed_immutable_and_internal_only():
    api = _api()
    job, _ = _job(api)

    assert tuple(signature(type(job)).parameters) == (
        "controlled_seed_store",
        "export_port",
        "office_qualification_port",
        "qualification_evidence_store",
    )
    assert tuple(signature(job.qualify).parameters) == (
        "document_type",
        "seed_handle",
        "fixture",
        "profile_id",
    )
    assert {item.name for item in fields(api.CandidateQualificationEvidence)} == {
        "qualification_id",
        "document_type",
        "seed_sha256",
        "profile_id",
        "profile_version",
        "rendered_sha256",
        "parse_report_sha256",
        "office_evidence_id",
        "office_client_versions",
        "fixture_id",
        "checker_version",
        "qualified_at_utc",
        "qualification_status",
    }
    fixture = _fixture(api)
    with pytest.raises(FrozenInstanceError):
        fixture.fixture_id = "changed"
    with pytest.raises(ValueError):
        api.SyntheticQualificationFixture(
            fixture_id="weekly-monthly-fixture-v1",
            provenance="synthetic",
            values=(("title", {"mutable": "value"}),),
        )


@pytest.mark.asyncio
@pytest.mark.parametrize(
    ("document_type", "seed_handle"),
    [
        ("weekly_activity_plan", "controlled-weekplan-seed-v2"),
        ("monthly_theme_activity_plan", "controlled-monthplan-seed-v2"),
    ],
)
async def test_reserved_candidate_qualification_uses_controlled_seed_and_same_opaque_export_port(
    document_type, seed_handle
):
    api = _api()
    job, effects = _job(api)

    evidence = await job.qualify(
        document_type=document_type,
        seed_handle=seed_handle,
        fixture=_fixture(api),
        profile_id=f"{document_type}-profile-v2",
    )

    assert evidence.document_type == document_type
    assert evidence.qualification_status == "passed"
    assert len(evidence.seed_sha256) == 64
    assert evidence.profile_id == f"{document_type}-profile-v2"
    assert len(evidence.rendered_sha256) == 64
    assert len(evidence.parse_report_sha256) == 64
    assert evidence.profile_version == 2
    assert evidence.office_evidence_id == "office-qualification-v1"
    assert evidence.office_client_versions
    assert evidence.fixture_id == "weekly-monthly-fixture-v1"
    assert type(evidence.qualification_id) is UUID
    assert evidence.checker_version == "template-candidate-qualification.v1"
    assert evidence.qualified_at_utc.utcoffset().total_seconds() == 0
    assert not hasattr(evidence, "rendered_bytes")
    assert not hasattr(evidence, "seed_bytes")
    assert not hasattr(evidence, "absolute_path")
    assert not hasattr(evidence, "active_version_id")
    assert not hasattr(evidence, "template_version_id")
    assert effects["seeds"].read_calls == [(seed_handle, document_type)]
    assert len(effects["export"].render_calls) == 1
    assert len(effects["export"].parse_calls) == 1
    binding, payload = effects["export"].render_calls[0]
    assert type(binding) is api.TemplateExportBinding
    assert binding.kind.value == "candidate"
    assert binding.document_type.value == document_type
    assert binding.content_sha256 == evidence.seed_sha256
    assert binding.profile_id == evidence.profile_id
    assert binding.profile_version == evidence.profile_version
    assert binding.tenant_id is None
    assert binding.template_version_id is None
    assert binding.version is None
    assert payload == _fixture(api).values
    parse_binding, rendered_bytes = effects["export"].parse_calls[0]
    assert parse_binding is binding
    assert sha256(rendered_bytes).hexdigest() == evidence.rendered_sha256
    office_binding, parse_report, office_profile = effects["office"].calls[0]
    assert office_binding is binding
    assert parse_report.binding is binding
    assert office_profile == evidence.profile_id
    assert effects["export"].resolve_calls == []
    assert len(effects["office"].calls) == 1
    assert effects["evidence"].items == [evidence]


@pytest.mark.asyncio
@pytest.mark.parametrize(
    ("document_type", "seed_handle", "profile_id"),
    [
        (
            "weekly_activity_plan",
            "controlled-monthplan-seed-v2",
            "weekly_activity_plan-profile-v2",
        ),
        (
            "monthly_theme_activity_plan",
            "controlled-weekplan-seed-v2",
            "monthly_theme_activity_plan-profile-v2",
        ),
        (
            "weekly_activity_plan",
            "controlled-weekplan-seed-v2",
            "monthly_theme_activity_plan-profile-v2",
        ),
        (
            "weekly_activity_plan",
            "controlled-weekplan-seed-v2",
            "candidate-profile-v999",
        ),
    ],
    ids=["cross-seed-weekly", "cross-seed-monthly", "cross-profile", "unknown-profile"],
)
async def test_candidate_preflight_closes_seed_and_profile_to_document_type(
    document_type, seed_handle, profile_id
):
    api = _api()
    job, effects = _job(api)

    with pytest.raises(api.TemplateCenterError) as caught:
        await job.qualify(
            document_type=document_type,
            seed_handle=seed_handle,
            fixture=_fixture(api),
            profile_id=profile_id,
        )

    assert caught.value.code is api.TemplateErrorCode.INPUT_INVALID
    assert seed_handle not in repr(caught.value)
    assert effects["seeds"].read_calls == []
    assert effects["export"].render_calls == []
    assert effects["office"].calls == []
    assert effects["evidence"].items == []


@pytest.mark.asyncio
@pytest.mark.parametrize(
    ("case_id", "seed_content"),
    [
        ("macro", docx_with_macro_member()),
        ("external-rel", docx_with_external_relationship()),
        ("bad-zip", b"registered candidate but not a ZIP package"),
        ("structure-profile-mismatch", docx_with_structure_profile_mismatch()),
    ],
    ids=["macro", "external-rel", "bad-zip", "structure-profile-mismatch"],
)
async def test_registered_candidate_seed_reuses_security_validator_and_fails_before_render_or_office(
    case_id, seed_content
):
    """A registered handle is not a trust bypass: security/profile failures stop the pipeline."""
    api = _api()
    seed_handle = "controlled-weekplan-seed-v2"
    seeds = MemoryControlledSeedStore()
    seeds.register_controlled_seed(seed_handle, seed_content)
    job, effects = _job(api, seeds=seeds)

    with pytest.raises(api.TemplateCenterError):
        await job.qualify(
            document_type="weekly_activity_plan",
            seed_handle=seed_handle,
            fixture=_fixture(api),
            profile_id="weekly_activity_plan-profile-v2",
        )

    # This is deliberately distinct from preflight rejection: the controlled seed
    # was read, but the shared byte/ZIP/XML/profile validator stopped the pipeline.
    assert effects["seeds"].read_calls == [(seed_handle, "weekly_activity_plan")]
    assert effects["export"].render_calls == []
    assert effects["export"].parse_calls == []
    assert effects["office"].calls == []
    assert not any(
        getattr(item, "qualification_status", None) == "passed"
        for item in effects["evidence"].items
    )
    assert effects["export"].resolve_calls == []
    assert (
        api.build_initial_document_registry().is_enabled("weekly_activity_plan")
        is False
    )


@pytest.mark.asyncio
@pytest.mark.parametrize(
    ("case_id", "status", "evidence_id", "client_versions"),
    [
        ("office-failed", "failed", "office-failure-v1", OFFICE_CLIENT_VERSIONS),
        ("missing-word", "passed", "office-without-word-v1", ("libreoffice/24.2.7.2",)),
        (
            "missing-libreoffice",
            "passed",
            "office-without-libreoffice-v1",
            ("word/16.0.17328.20124",),
        ),
        ("missing-evidence-id", "passed", None, OFFICE_CLIENT_VERSIONS),
        (
            "missing-exact-client-version",
            "passed",
            "office-unversioned-v1",
            ("word", "libreoffice"),
        ),
    ],
    ids=[
        "office-failed",
        "missing-word",
        "missing-libreoffice",
        "missing-evidence-id",
        "missing-exact-client-version",
    ],
)
async def test_incomplete_office_qualification_fails_after_render_without_passed_evidence_or_enablement(
    case_id, status, evidence_id, client_versions
):
    """Office evidence is an independent required gate, not a best-effort annotation."""
    api = _api()
    office = MemoryOfficeQualificationPort(
        status=status,
        evidence_id=evidence_id,
        client_versions=client_versions,
    )
    job, effects = _job(api, office=office)

    with pytest.raises(api.TemplateCenterError):
        await job.qualify(
            document_type="weekly_activity_plan",
            seed_handle="controlled-weekplan-seed-v2",
            fixture=_fixture(api),
            profile_id="weekly_activity_plan-profile-v2",
        )

    # The seed passed the common security/profile stage, so render/parse and the
    # Office call are observable; the incomplete Office result must still fail closed.
    assert effects["seeds"].read_calls == [
        ("controlled-weekplan-seed-v2", "weekly_activity_plan")
    ]
    assert len(effects["export"].render_calls) == 1
    assert len(effects["export"].parse_calls) == 1
    assert len(effects["office"].calls) == 1
    assert not any(
        getattr(item, "qualification_status", None) == "passed"
        for item in effects["evidence"].items
    )
    assert effects["export"].resolve_calls == []
    assert (
        api.build_initial_document_registry().is_enabled("weekly_activity_plan")
        is False
    )


@pytest.mark.asyncio
@pytest.mark.parametrize(
    ("render_mode", "parse_mode", "office_raises", "expected_parse", "expected_office"),
    [
        ("raises", "valid", False, 0, 0),
        ("binding-mismatch", "valid", False, 0, 0),
        ("hash-mismatch", "valid", False, 0, 0),
        ("valid", "raises", False, 1, 0),
        ("valid", "binding-mismatch", False, 1, 0),
        ("valid", "invalid", False, 1, 0),
        ("valid", "unresolved", False, 1, 0),
        ("valid", "macro", False, 1, 0),
        ("valid", "external", False, 1, 0),
        ("valid", "valid", True, 1, 1),
    ],
    ids=[
        "render-error",
        "render-binding",
        "render-hash",
        "parse-error",
        "parse-binding",
        "parse-invalid",
        "parse-unresolved",
        "parse-macro",
        "parse-external",
        "office-error",
    ],
)
async def test_candidate_pipeline_failures_short_circuit_without_evidence(
    render_mode, parse_mode, office_raises, expected_parse, expected_office
):
    api = _api()
    export = MemoryExportPort(render_mode=render_mode, parse_mode=parse_mode)
    office = MemoryOfficeQualificationPort(raises=office_raises)
    job, effects = _job(api, export=export, office=office)

    with pytest.raises(api.TemplateCenterError) as caught:
        await job.qualify(
            document_type="weekly_activity_plan",
            seed_handle="controlled-weekplan-seed-v2",
            fixture=_fixture(api),
            profile_id="weekly_activity_plan-profile-v2",
        )

    assert caught.value.code is api.TemplateErrorCode.EXPORT_FAILED
    assert len(effects["export"].parse_calls) == expected_parse
    assert len(effects["office"].calls) == expected_office
    assert effects["evidence"].items == []
    assert effects["export"].resolve_calls == []


@pytest.mark.asyncio
@pytest.mark.parametrize(
    ("render_mode", "parse_mode", "expected_parse"),
    [
        ("unsafe-bytes", "valid", 0),
        ("valid", "structure-mismatch", 1),
    ],
    ids=["rendered-not-docx", "parse-structure-not-bound"],
)
async def test_render_and_parse_cannot_self_attest_unbound_content(
    render_mode, parse_mode, expected_parse
):
    api = _api()
    export = MemoryExportPort(render_mode=render_mode, parse_mode=parse_mode)
    job, effects = _job(api, export=export)

    with pytest.raises(api.TemplateCenterError) as caught:
        await job.qualify(
            document_type="weekly_activity_plan",
            seed_handle="controlled-weekplan-seed-v2",
            fixture=_fixture(api),
            profile_id="weekly_activity_plan-profile-v2",
        )

    assert caught.value.code is api.TemplateErrorCode.EXPORT_FAILED
    assert len(effects["export"].parse_calls) == expected_parse
    assert effects["office"].calls == []
    assert effects["evidence"].items == []


@pytest.mark.asyncio
async def test_candidate_evidence_append_failure_is_sanitized_and_not_visible():
    api = _api()
    evidence_store = MemoryQualificationEvidenceStore(fail=True)
    job, effects = _job(api, evidence=evidence_store)

    with pytest.raises(api.TemplateCenterError) as caught:
        await job.qualify(
            document_type="weekly_activity_plan",
            seed_handle="controlled-weekplan-seed-v2",
            fixture=_fixture(api),
            profile_id="weekly_activity_plan-profile-v2",
        )

    assert caught.value.code is api.TemplateErrorCode.STORAGE_FAILED
    assert effects["evidence"].items == []
    assert "synthetic" not in repr(caught.value)


@pytest.mark.asyncio
@pytest.mark.parametrize(
    ("document_type", "seed_handle", "provenance"),
    [
        ("weekly_activity_plan", "../templates/weekplan.docx", "synthetic"),
        ("monthly_theme_activity_plan", "controlled-monthplan-seed-v2", "business"),
        ("daily_plan", "controlled-weekplan-seed-v2", "synthetic"),
        ("unknown_type", "controlled-weekplan-seed-v2", "synthetic"),
    ],
)
async def test_candidate_qualification_rejects_uncontrolled_or_non_reserved_inputs_without_side_effects(
    document_type, seed_handle, provenance
):
    api = _api()
    job, effects = _job(api)

    with pytest.raises(api.TemplateCenterError) as caught:
        await job.qualify(
            document_type=document_type,
            seed_handle=seed_handle,
            fixture=_fixture(api, provenance=provenance),
            profile_id="candidate-profile-v1",
        )

    assert seed_handle not in repr(caught.value)
    assert effects["seeds"].read_calls == []
    assert effects["export"].render_calls == []
    assert effects["export"].parse_calls == []
    assert effects["export"].resolve_calls == []
    assert effects["office"].calls == []
    assert effects["evidence"].items == []


@pytest.mark.asyncio
async def test_qualification_is_not_formal_preview_or_enablement_and_exposes_no_public_crud():
    api = _api()
    job, effects = _job(api)

    evidence = await job.qualify(
        document_type=RESERVED_TYPES[0],
        seed_handle="controlled-weekplan-seed-v2",
        fixture=_fixture(api),
        profile_id="weekly_activity_plan-profile-v2",
    )
    registry = api.build_initial_document_registry()

    assert evidence.qualification_status == "passed"
    assert registry.is_enabled(RESERVED_TYPES[0]) is False
    assert tuple(descriptor.key for descriptor in registry.descriptors()) == (
        "daily_plan",
        "game_observation",
        "one_on_one_listening",
        "homemade_teaching",
        "course_review_activity",
    )
    for forbidden in (
        "project",
        "upload",
        "preview",
        "resolve_active",
        "activate",
        "deactivate",
        "rollback",
        "download",
        "formal_download",
        "read_business",
        "write_business",
        "create_version",
    ):
        assert not hasattr(job, forbidden)
    assert not hasattr(evidence, "rendered_bytes")
    assert not hasattr(evidence, "content")
    assert not hasattr(evidence, "active_version_id")
    assert effects["export"].resolve_calls == []
