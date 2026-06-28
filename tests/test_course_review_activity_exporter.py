import io
from pathlib import Path

from docx import Document

from app.integration.word_export.course_review_activity_exporter import (
    export_course_review_activity,
)


def _make_record(**kwargs) -> dict:
    data = {
        "grade": "小班",
        "class_name": "阳光班",
        "teacher_name": "张老师",
        "activity_name": "彩虹桥搭建",
        "child_count": "28",
        "activity_time": "2026.06.28",
        "lesson_plan_original": "原始教案：彩虹桥搭建\n活动目标：尝试搭建。",
        "activity_goal": "1.尝试搭建彩虹桥。",
        "activity_prep": "积木、彩色纸条。",
        "activity_process": "教师示范，幼儿搭建。",
        "goal_adjusted": False,
        "goal_adjustment": "",
        "activity_goal_revised": "1.尝试搭建彩虹桥。",
        "prep_adjusted": True,
        "prep_adjustment": "增加低结构辅助材料。",
        "activity_prep_revised": "积木、彩色纸条、纸筒。",
        "process_adjustment": "增加幼儿自由探索和同伴分享环节。",
        "activity_process_revised": "幼儿先自由探索，再合作搭建。",
        "review_reason": "更符合小班幼儿先探索再表达的学习特点。",
        "revised_lesson_plan": "二次修改稿：彩虹桥搭建\n活动目标：尝试搭建彩虹桥。",
    }
    data.update(kwargs)
    return data


def _parse(doc_bytes: bytes) -> Document:
    return Document(io.BytesIO(doc_bytes))


def test_export_course_review_activity_returns_parseable_docx():
    result = export_course_review_activity(_make_record())

    assert isinstance(result, bytes)
    assert len(result) > 0
    doc = _parse(result)
    assert len(doc.tables) == 1
    assert len(doc.tables[0].rows) == 8


def test_export_course_review_activity_fills_template_cells():
    doc = _parse(export_course_review_activity(_make_record()))
    table = doc.tables[0]

    assert table.rows[0].cells[1].text == "彩虹桥搭建"
    assert table.rows[1].cells[3].text == "小班"
    assert table.rows[1].cells[5].text == "28"
    assert table.rows[2].cells[1].text == "张老师"
    assert table.rows[2].cells[3].text == "张老师"
    assert table.rows[2].cells[5].text == "2026.06.28"
    assert "尝试搭建彩虹桥" in table.rows[4].cells[1].text
    assert "积木" in table.rows[5].cells[1].text
    assert "教师示范" in table.rows[6].cells[1].text
    assert "小班幼儿" in table.rows[7].cells[1].text


def test_export_course_review_activity_writes_adjustment_marks():
    doc = _parse(export_course_review_activity(_make_record()))
    table = doc.tables[0]

    goal_revision = table.rows[4].cells[4].text
    prep_revision = table.rows[5].cells[4].text
    process_revision = table.rows[6].cells[4].text

    assert "保持原设计不变（√）" in goal_revision
    assert "有所调整（）" in goal_revision
    assert "保持原设计，不变（）" in prep_revision
    assert "有所调整（√）" in prep_revision
    assert "低结构辅助材料" in prep_revision
    assert process_revision.startswith("调整内容：")
    assert "同伴分享" in process_revision


def test_export_course_review_activity_writes_original_and_revised_drafts():
    doc = _parse(export_course_review_activity(_make_record()))
    full_text = "\n".join(p.text for p in doc.paragraphs)

    assert "附教案：原稿" in full_text
    assert "原始教案：彩虹桥搭建" in full_text
    assert "二次修改" in full_text
    assert "二次修改稿：彩虹桥搭建" in full_text


def test_export_course_review_activity_removes_template_sample_content():
    doc = _parse(export_course_review_activity(_make_record()))
    table_text = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                table_text.append(cell.text)
    full_text = "\n".join([p.text for p in doc.paragraphs] + table_text)

    assert "贝壳社区" not in full_text
    assert "圆形灯笼" not in full_text
    assert len(doc.tables) == 1


def test_export_course_review_activity_fallback_when_template_missing(tmp_path: Path):
    doc = _parse(
        export_course_review_activity(
            _make_record(),
            template_path=tmp_path / "missing.docx",
        )
    )

    assert len(doc.tables) == 1
    assert len(doc.tables[0].rows) == 8
    assert "彩虹桥搭建" in doc.tables[0].rows[0].cells[1].text
    assert "二次修改稿" in "\n".join(p.text for p in doc.paragraphs)
