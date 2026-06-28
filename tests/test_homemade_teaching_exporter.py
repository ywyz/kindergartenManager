import io
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from app.integration.word_export.homemade_teaching_exporter import (
    export_homemade_teaching,
)


def _make_record(**kwargs) -> dict:
    data = {
        "class_name": "阳光班",
        "teacher_name": "张老师",
        "toy_name": "彩虹穿线板",
        "materials": "硬纸板、毛根、打孔器",
        "play_methods": "幼儿按颜色规律穿线。",
    }
    data.update(kwargs)
    return data


def _parse(doc_bytes: bytes) -> Document:
    return Document(io.BytesIO(doc_bytes))


def test_export_homemade_teaching_returns_parseable_docx():
    result = export_homemade_teaching(_make_record())

    assert isinstance(result, bytes)
    assert len(result) > 0
    doc = _parse(result)
    assert len(doc.tables) == 1
    assert len(doc.tables[0].rows) == 5


def test_export_homemade_teaching_fills_template_cells():
    result = export_homemade_teaching(
        _make_record(
            class_name="彩虹班",
            teacher_name="李老师",
            toy_name="瓶盖配对盒",
            materials="瓶盖、纸盒、彩色贴纸",
            play_methods="幼儿根据颜色和大小进行配对。",
        )
    )
    doc = _parse(result)
    table = doc.tables[0]

    assert table.rows[0].cells[1].text == "彩虹班"
    assert table.rows[1].cells[1].text == "李老师"
    assert table.rows[2].cells[1].text == "瓶盖配对盒"
    assert "彩色贴纸" in table.rows[3].cells[1].text
    assert "配对" in table.rows[4].cells[1].text


def test_export_homemade_teaching_sets_chinese_font():
    result = export_homemade_teaching(_make_record(toy_name="字体测试"))
    doc = _parse(result)
    cell = doc.tables[0].rows[2].cells[1]

    found_songti = False
    for para in cell.paragraphs:
        for run in para.runs:
            if run.text.strip():
                rpr = run._element.find(qn("w:rPr"))
                if rpr is None:
                    continue
                r_fonts = rpr.find(qn("w:rFonts"))
                if r_fonts is not None and r_fonts.get(qn("w:eastAsia")) == "宋体":
                    found_songti = True

    assert found_songti


def test_export_homemade_teaching_fallback_when_template_missing(tmp_path: Path):
    result = export_homemade_teaching(
        _make_record(),
        template_path=tmp_path / "missing.docx",
    )

    doc = _parse(result)
    assert len(doc.tables) == 1
    assert len(doc.tables[0].rows) == 5
    assert "彩虹穿线板" in doc.tables[0].rows[2].cells[1].text
