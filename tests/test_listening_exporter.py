"""P6 — 一对一倾听 Word 导出器测试（真实模板）。"""
import io
from datetime import date
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn

from app.integration.word_export.listening_exporter import (
    DOMAINS,
    export_batch_by_domain,
    export_combined,
    export_split_by_domain,
    _norm,
)


def _tiny_jpeg() -> bytes:
    from PIL import Image
    buf = io.BytesIO()
    Image.new("RGB", (40, 30), (180, 100, 60)).save(buf, format="JPEG")
    return buf.getvalue()


_IMG = _tiny_jpeg()
_RECORD = {"child_name": "小明", "adult_count": 1, "child_age": "4岁"}
_IND_COUNT = {"健康": 7, "语言": 6, "社会": 7, "科学": 6, "艺术": 4}


def _domain_payload(domain: str, star_for_idx1: int = 2) -> dict:
    n = _IND_COUNT[domain]
    return {
        "domain": domain, "obs_year": 2026, "obs_month": 4,
        "date_1": date(2026, 4, 1), "date_2": date(2026, 4, 6), "date_3": date(2026, 4, 13),
        "goals": f"{domain}目标A；{domain}目标B",
        "evaluation": f"{domain}综合评价正文", "support_strategy": f"{domain}支持策略正文",
        "images": [(_IMG, f"{domain}图{i+1}描述") for i in range(3)],
        "indicators": [
            {"sort_order": i, "stars": (star_for_idx1 if i == 1 else 3)} for i in range(n)
        ],
    }


def _all_domains() -> list[dict]:
    return [_domain_payload(d) for d in DOMAINS]


# ─── 合并导出 ────────────────────────────────────────────────────────────────


def test_combined_reopen_has_five_tables():
    doc = Document(io.BytesIO(export_combined(_RECORD, _all_domains())))
    assert len(doc.tables) == 5


def test_metadata_filled():
    doc = Document(io.BytesIO(export_combined(_RECORD, _all_domains())))
    meta = doc.tables[0].rows[0].cells[0].text
    assert "2026年4月" in meta
    assert "小明" in meta
    assert "4岁" in meta
    assert "健康目标A" in meta


def test_dates_in_drawing_headers():
    doc = Document(io.BytesIO(export_combined(_RECORD, _all_domains())))
    t0 = doc.tables[0]
    assert "4月1日" in t0.rows[1].cells[0].text
    assert "4月6日" in t0.rows[3].cells[0].text
    assert "4月13日" in t0.rows[5].cells[0].text


def test_indicator_check_mark_position():
    """健康 sort_order=1 指标评 2 星 → R12 的 C5 为 √，R11/R13 为空。"""
    doc = Document(io.BytesIO(export_combined(_RECORD, _all_domains())))
    t0 = doc.tables[0]
    assert t0.rows[12].cells[5].text.strip() == "√"
    assert t0.rows[11].cells[5].text.strip() == ""
    assert t0.rows[13].cells[5].text.strip() == ""


def test_check_mark_three_stars_default():
    """默认 3 星指标（sort_order=0）→ ★★★ 行（R10）C5 为 √。"""
    doc = Document(io.BytesIO(export_combined(_RECORD, _all_domains())))
    t0 = doc.tables[0]
    # 指标0：行 R8/R9/R10 → 3 星标记 R10
    assert t0.rows[10].cells[5].text.strip() == "√"
    assert t0.rows[8].cells[5].text.strip() == ""


def test_images_inserted_combined():
    doc = Document(io.BytesIO(export_combined(_RECORD, _all_domains())))
    assert len(doc.inline_shapes) == 15  # 5 领域 × 3 张


def test_eval_and_strategy_cells():
    doc = Document(io.BytesIO(export_combined(_RECORD, _all_domains())))
    t0 = doc.tables[0]
    zh = next(i for i, r in enumerate(t0.rows) if _norm(r.cells[0].text) == "综合评价")
    st = next(i for i, r in enumerate(t0.rows) if _norm(r.cells[0].text) == "支持策略")
    assert "健康综合评价正文" in t0.rows[zh].cells[1].text
    assert "健康支持策略正文" in t0.rows[st].cells[1].text


def test_chinese_font_song():
    """写入的 run 字体为宋体且设置 eastAsia。"""
    doc = Document(io.BytesIO(export_combined(_RECORD, _all_domains())))
    t0 = doc.tables[0]
    zh = next(i for i, r in enumerate(t0.rows) if _norm(r.cells[0].text) == "综合评价")
    run = t0.rows[zh].cells[1].paragraphs[0].runs[0]
    assert run.font.name == "宋体"
    rfonts = run._element.rPr.rFonts
    assert rfonts.get(qn("w:eastAsia")) == "宋体"


# ─── 拆分导出 ────────────────────────────────────────────────────────────────


def test_split_by_domain_keys():
    result = export_split_by_domain(_RECORD, _all_domains())
    assert set(result.keys()) == set(DOMAINS)
    for d, b in result.items():
        sd = Document(io.BytesIO(b))
        assert len(sd.tables) == 1
        assert len(sd.inline_shapes) == 3


def test_split_only_domains_with_data():
    """仅含数据的领域才导出。"""
    result = export_split_by_domain(_RECORD, [_domain_payload("健康"), _domain_payload("语言")])
    assert set(result.keys()) == {"健康", "语言"}


# ─── 批量导出 ────────────────────────────────────────────────────────────────


def test_batch_by_domain_two_children():
    record2 = {"child_name": "小红", "adult_count": 1, "child_age": "4岁"}
    result = export_batch_by_domain([(_RECORD, _all_domains()), (record2, _all_domains())])
    assert set(result.keys()) == set(DOMAINS)
    hb = Document(io.BytesIO(result["健康"]))
    assert len(hb.tables) == 2  # 2 幼儿
    assert len(hb.inline_shapes) == 6  # 2 × 3
    allmeta = " ".join(t.rows[0].cells[0].text for t in hb.tables)
    assert "小明" in allmeta and "小红" in allmeta


def test_batch_custom_domain_order():
    order = ["语言", "艺术", "健康"]
    result = export_batch_by_domain([(_RECORD, _all_domains())], domain_order=order)
    assert list(result.keys()) == order


# ─── 模板缺失降级 ────────────────────────────────────────────────────────────


def test_template_missing_fallback_combined():
    data = export_combined(_RECORD, _all_domains(), template_path=Path("/nonexistent/x.docx"))
    doc = Document(io.BytesIO(data))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "健康" in text
    assert "小明" in text


def test_template_missing_fallback_split():
    result = export_split_by_domain(_RECORD, [_domain_payload("健康")], template_path=Path("/no/x.docx"))
    assert "健康" in result
    Document(io.BytesIO(result["健康"]))  # 可打开，不抛异常
