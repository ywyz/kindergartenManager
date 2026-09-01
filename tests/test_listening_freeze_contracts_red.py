"""Stable RED for the frozen one-on-one listening export contract."""

from __future__ import annotations

import inspect
import io
from datetime import date

from docx import Document
from docx.oxml.ns import qn

from app.integration.word_export import listening_exporter
from app.ui.pages import one_on_one_listening

FROZEN_DOMAINS = ["健康", "语言", "社会", "科学", "艺术"]


def _domain(domain: str) -> dict:
    return {
        "domain": domain,
        "obs_year": 2026,
        "obs_month": 4,
        "date_1": date(2026, 4, 1),
        "date_2": date(2026, 4, 2),
        "date_3": date(2026, 4, 3),
        "goals": f"{domain}目标",
        "evaluation": f"{domain}评价",
        "support_strategy": f"{domain}支持",
        "images": [],
        "indicators": [],
    }


def _record(name: str) -> dict:
    return {"child_name": name, "adult_count": 1, "child_age": "4岁"}


def _table_names(doc: Document) -> list[str]:
    return [table.rows[0].cells[0].text for table in doc.tables]


def test_domain_order_is_frozen_across_exporter_and_ui() -> None:
    assert listening_exporter.DOMAINS == FROZEN_DOMAINS
    assert one_on_one_listening._UI_DOMAINS == FROZEN_DOMAINS


def test_partial_combined_removes_unselected_template_blocks() -> None:
    data = listening_exporter.export_combined(
        _record("小明"),
        [_domain("艺术"), _domain("健康")],
    )
    doc = Document(io.BytesIO(data))

    assert len(doc.tables) == 2
    assert "健康目标" in doc.tables[0].rows[0].cells[0].text
    assert "艺术目标" in doc.tables[1].rows[0].cells[0].text
    body_text = "".join(node.text or "" for node in doc.element.body.iter(qn("w:t")))
    assert "语言领域" not in body_text
    assert "社会领域" not in body_text
    assert "科学领域" not in body_text


def test_batch_combined_preserves_child_selection_and_template_domain_order() -> None:
    export_batch_combined = getattr(listening_exporter, "export_batch_combined", None)
    assert callable(export_batch_combined)
    children = [
        (_record("第二个"), [_domain("科学"), _domain("健康")]),
        (_record("第一个"), [_domain("健康"), _domain("科学")]),
    ]

    doc = Document(io.BytesIO(export_batch_combined(children)))

    assert len(doc.tables) == 4
    names = _table_names(doc)
    assert "第二个" in names[0] and "健康目标" in names[0]
    assert "第二个" in names[1] and "科学目标" in names[1]
    assert "第一个" in names[2] and "健康目标" in names[2]
    assert "第一个" in names[3] and "科学目标" in names[3]
    page_breaks = doc.element.body.xpath('.//w:br[@w:type="page"]')
    assert len(page_breaks) == 1


def test_batch_by_child_emits_one_filtered_document_per_child_in_order() -> None:
    export_batch_by_child = getattr(listening_exporter, "export_batch_by_child", None)
    assert callable(export_batch_by_child)
    children = [
        (_record("乙"), [_domain("语言")]),
        (_record("甲"), [_domain("健康")]),
    ]

    files = export_batch_by_child(children)

    assert [name for name, _data in files] == ["乙", "甲"]
    for _name, data in files:
        assert len(Document(io.BytesIO(data)).tables) == 1


def test_ordered_selection_helper_never_sorts_user_choice() -> None:
    update = getattr(one_on_one_listening, "ordered_selection_after", None)
    assert callable(update)

    selected = update((), 9, True)
    selected = update(selected, 2, True)
    selected = update(selected, 7, True)
    selected = update(selected, 2, False)
    selected = update(selected, 2, True)

    assert selected == (9, 7, 2)


def test_page_exposes_domain_selection_and_all_three_batch_modes() -> None:
    source = inspect.getsource(one_on_one_listening.one_on_one_listening_page)

    assert "选择导出领域" in source
    assert "批量合并 Word" in source
    assert "批量 ZIP（按幼儿）" in source
    assert "批量 ZIP（按领域）" in source
