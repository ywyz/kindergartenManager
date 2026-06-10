"""tests/test_observation_exporter.py — 游戏观察 Word 导出测试。

测试覆盖（8项）：
  1. 导出返回非空 bytes，可被 python-docx 重新打开
  2. 标题中 'xx' 被替换为大环境值
  3. 各关键单元格内容正确写入
  4. 模板示例文本被清空（不残留「王鹤宁、侯舒妍」）
  5. 传入 1 张图片时，R5 观察记录单元格内含 1 张图片
  6. 传入 3 张图片时，R5 观察记录单元格内含 3 张图片
  7. 中文 run 字体为宋体，设置了 w:eastAsia
  8. 模板缺失时降级从零构表，不报错，返回可打开的 bytes
"""
import io
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn

from app.integration.word_export.observation_exporter import export_observation


def _make_obs(**kwargs) -> dict:
    """构造测试用观察记录字典。"""
    defaults = {
        "class_name": "阳光班",
        "obs_date": "2026-06-10",
        "time_range": "9:00-9:40",
        "big_env": "户外",
        "game_area": "建构区",
        "adult_count": 2,
        "child_count": 15,
        "child_names": "小明、小红",
        "child_age": "5岁",
        "observer": "李老师",
        "observation_goal": "观察幼儿合作行为",
        "observation_record": "幼儿积极参与搭建活动",
        "evaluation_analysis": "表现出良好的合作意识",
        "support_strategy": "适时提问引导协商",
    }
    defaults.update(kwargs)
    return defaults


_TINY_JPEG = (
    b"\xff\xd8\xff\xe0\x00\x10JFIF\x00\x01\x01\x00\x00\x01\x00\x01\x00\x00"
    b"\xff\xdb\x00C\x00\x08\x06\x06\x07\x06\x05\x08\x07\x07\x07\t\t"
    b"\x08\n\x0c\x14\r\x0c\x0b\x0b\x0c\x19\x12\x13\x0f\x14\x1d\x1a"
    b"\x1f\x1e\x1d\x1a\x1c\x1c $.' \",#\x1c\x1c(7),01444\x1f'9=82<.342\x1e"
    b"A+*\xa2\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00"
    b"\xff\xc0\x00\x0b\x08\x00\x01\x00\x01\x01\x01\x11\x00\xff\xc4\x00\x1f\x00\x00\x01"
    b"\x05\x01\x01\x01\x01\x01\x01\x00\x00\x00\x00\x00\x00\x00\x00\x01\x02\x03\x04\x05"
    b"\x06\x07\x08\t\n\x0b\xff\xda\x00\x08\x01\x01\x00\x00?\x00\xf5\x7f\xff\xd9"
)


def _parse(doc_bytes: bytes) -> Document:
    return Document(io.BytesIO(doc_bytes))


# ---------------------------------------------------------------------------
# 1. 返回非空 bytes，可被 python-docx 打开
# ---------------------------------------------------------------------------


def test_export_returns_nonempty_parseable_bytes():
    """导出返回非空 bytes，且 python-docx 可成功解析。"""
    result = export_observation(_make_obs(), [])
    assert isinstance(result, bytes)
    assert len(result) > 0
    doc = _parse(result)
    assert len(doc.tables) >= 1


# ---------------------------------------------------------------------------
# 2. 标题中 'xx' 被替换为大环境
# ---------------------------------------------------------------------------


def test_title_xx_replaced_by_big_env():
    """标题段落中的 'xx' 应被替换为大环境值（如'户外'）。"""
    obs = _make_obs(big_env="室内")
    result = export_observation(obs, [])
    doc = _parse(result)
    title_text = doc.paragraphs[0].text
    assert "xx" not in title_text
    assert "室内" in title_text


# ---------------------------------------------------------------------------
# 3. 关键单元格内容正确写入
# ---------------------------------------------------------------------------


def test_key_cell_contents():
    """验证班级、观察目标、观察记录、支持策略等关键字段写入正确。"""
    obs = _make_obs(
        class_name="大班",
        observation_goal="测试观察目标",
        observation_record="测试观察记录",
        support_strategy="测试支持策略",
    )
    result = export_observation(obs, [])
    doc = _parse(result)
    full_text = "\n".join(
        c.text for row in doc.tables[0].rows for c in row.cells
    )
    assert "大班" in full_text
    assert "测试观察目标" in full_text
    assert "测试观察记录" in full_text
    assert "测试支持策略" in full_text


# ---------------------------------------------------------------------------
# 4. 示例文本被清空
# ---------------------------------------------------------------------------


def test_example_text_cleared():
    """模板示例文本「王鹤宁、侯舒妍」应被清空，不残留在导出文件中。"""
    obs = _make_obs(child_names="新名字")
    result = export_observation(obs, [])
    doc = _parse(result)
    full_text = "\n".join(
        c.text for row in doc.tables[0].rows for c in row.cells
    )
    assert "王鹤宁" not in full_text
    assert "侯舒妍" not in full_text


# ---------------------------------------------------------------------------
# 5. 1 张图片 → R5 含 1 张图片
# ---------------------------------------------------------------------------


def test_one_image_in_r5():
    """传入 1 张图片时，R5 观察记录单元格应包含 1 个图片元素。"""
    result = export_observation(_make_obs(), [_TINY_JPEG])
    doc = _parse(result)
    row5_cell = doc.tables[0].rows[5].cells[1]
    # 统计所有 <a:blip> 图片引用（实际插入图片后会有此元素）
    blips = row5_cell._element.findall(
        ".//{http://schemas.openxmlformats.org/drawingml/2006/spreadsheetDrawing}blip"
    ) or row5_cell._element.findall(
        ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
    )
    # 用 drawing 元素数量作为图片数量（每个图片对应一个 drawing）
    drawings = row5_cell._element.findall(
        ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
    ) or row5_cell._element.findall(
        ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict"
    )
    # 兼容：用 inline/anchor 元素判断
    ns_wp = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    inline_elems = row5_cell._element.findall(f".//{{{ns_wp}}}inline")
    anchor_elems = row5_cell._element.findall(f".//{{{ns_wp}}}anchor")
    image_count = len(inline_elems) + len(anchor_elems)
    assert image_count == 1, f"期望 1 张图片，实际 {image_count} 张"


# ---------------------------------------------------------------------------
# 6. 3 张图片 → R5 含 3 张图片
# ---------------------------------------------------------------------------


def test_three_images_in_r5():
    """传入 3 张图片时，R5 观察记录单元格应包含 3 个图片元素。"""
    result = export_observation(_make_obs(), [_TINY_JPEG, _TINY_JPEG, _TINY_JPEG])
    doc = _parse(result)
    row5_cell = doc.tables[0].rows[5].cells[1]
    ns_wp = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    inline_elems = row5_cell._element.findall(f".//{{{ns_wp}}}inline")
    anchor_elems = row5_cell._element.findall(f".//{{{ns_wp}}}anchor")
    image_count = len(inline_elems) + len(anchor_elems)
    assert image_count == 3, f"期望 3 张图片，实际 {image_count} 张"


# ---------------------------------------------------------------------------
# 7. 中文 run 字体为宋体
# ---------------------------------------------------------------------------


def test_chinese_font_is_songti():
    """观察目标单元格中 run 的东亚字体应为宋体。"""
    obs = _make_obs(observation_goal="宋体字体测试")
    result = export_observation(obs, [])
    doc = _parse(result)
    row4_cell = doc.tables[0].rows[4].cells[1]
    # 找到所有 run 中有文字的
    found_songti = False
    for para in row4_cell.paragraphs:
        for run in para.runs:
            if run.text.strip():
                rpr = run._element.find(qn("w:rPr"))
                if rpr is not None:
                    r_fonts = rpr.find(qn("w:rFonts"))
                    if r_fonts is not None:
                        east_asia = r_fonts.get(qn("w:eastAsia"))
                        if east_asia == "宋体":
                            found_songti = True
    assert found_songti, "未发现宋体 w:eastAsia 字体设置"


# ---------------------------------------------------------------------------
# 8. 模板缺失 → 降级从零构表
# ---------------------------------------------------------------------------


def test_fallback_when_template_missing(tmp_path, monkeypatch):
    """模板文件缺失时，降级从零构建不报错，返回可打开的 bytes。"""
    import app.integration.word_export.observation_exporter as mod
    monkeypatch.setattr(mod, "TEMPLATE_PATH", tmp_path / "nonexistent.docx")

    result = export_observation(_make_obs(), [])
    assert isinstance(result, bytes)
    assert len(result) > 0
    doc = _parse(result)
    assert len(doc.tables) >= 1
