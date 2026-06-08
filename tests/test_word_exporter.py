"""Word 导出服务单元测试。

使用 python-docx 解析导出的 bytes，验证：
- 返回非空 bytes
- 使用模板时表格共 19 行，子字段分落到各单元格
- 周次/日期文本正确写入
- diff 差异中 changed=True 的句子字体为红色（活动过程格 R11）
- 无 diff 时直接输出改写文
- 空字段也能生成合法文档
- 行结构 AI 文本能解析为子字段
- 模板缺失时降级为从零构建（8 行）
"""
from datetime import date
from io import BytesIO

import pytest
from docx import Document
from docx.shared import RGBColor

from app.core.models.daily_plan import DailyPlan
from app.integration.word_export import exporter as exporter_mod
from app.integration.word_export.exporter import (
    _export_from_scratch,
    _parse_fields,
    export_batch_daily_plans,
    export_daily_plan,
)


def _make_plan(**kwargs) -> DailyPlan:
    """构造测试用 DailyPlan 实例（不依赖数据库，通过 SQLAlchemy __init__）。"""
    defaults: dict = {
        "tenant_id": 1,
        "user_id": 1,
        "plan_date": date(2026, 5, 18),
        "week_number": 3,
        "weekday_cn": "周一",
        "grade": "中班",
        "class_name": "阳光班",
        "activity_goal": "培养幼儿合作能力。",
        "activity_prep": "彩色积木若干。",
        "activity_key": "学会分工合作。",
        "activity_difficult": "协调动作一致。",
        "activity_process_original": "幼儿观察。老师示范。幼儿练习。",
        "activity_process_adapted": "幼儿仔细观察。老师耐心示范。幼儿反复练习。",
        "morning_activity": "体能大循环：跳绳、爬梯。",
        "indoor_area": "美工区：拼贴画。",
        "outdoor_activity": "追逐跑游戏。",
        "morning_talk_topic": "春天的花朵",
        "morning_talk_questions": "你最喜欢哪种花？为什么？",
        "daily_reflection": None,
    }
    defaults.update(kwargs)
    return DailyPlan(**defaults)


def _parse(doc_bytes: bytes) -> Document:
    return Document(BytesIO(doc_bytes))


class TestExportDailyPlan:
    def test_returns_nonempty_bytes(self):
        plan = _make_plan()
        result = export_daily_plan(plan, [])
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_table_has_template_rows(self):
        plan = _make_plan()
        result = export_daily_plan(plan, [])
        doc = _parse(result)
        assert len(doc.tables) == 1
        assert len(doc.tables[0].rows) == 19

    def test_week_number_in_first_row(self):
        plan = _make_plan(week_number=5)
        result = export_daily_plan(plan, [])
        doc = _parse(result)
        row0_text = doc.tables[0].rows[0].cells[0].text
        assert "5" in row0_text

    def test_date_in_second_row(self):
        plan = _make_plan(plan_date=date(2026, 5, 18), weekday_cn="周一")
        result = export_daily_plan(plan, [])
        doc = _parse(result)
        row1_text = doc.tables[0].rows[1].cells[0].text
        assert "5" in row1_text
        assert "18" in row1_text
        assert "周一" in row1_text

    def test_diff_changed_run_is_red(self):
        diff_result = [
            {"text": "幼儿仔细观察。", "changed": True},
            {"text": "老师耐心示范。", "changed": False},
            {"text": "幼儿反复练习。", "changed": True},
        ]
        plan = _make_plan()
        result = export_daily_plan(plan, diff_result)
        doc = _parse(result)

        content_cell = doc.tables[0].rows[11].cells[1]
        red_found = any(
            run.font.color.rgb == RGBColor(0xFF, 0x00, 0x00)
            for para in content_cell.paragraphs
            for run in para.runs
        )
        assert red_found, "应有 changed=True 的 run 使用红色字体"

    def test_unchanged_run_is_not_red(self):
        diff_result = [
            {"text": "幼儿仔细观察。", "changed": True},
            {"text": "老师耐心示范。", "changed": False},
        ]
        plan = _make_plan()
        result = export_daily_plan(plan, diff_result)
        doc = _parse(result)

        content_cell = doc.tables[0].rows[11].cells[1]
        for para in content_cell.paragraphs:
            for run in para.runs:
                if "老师耐心示范" in run.text:
                    assert run.font.color.rgb != RGBColor(0xFF, 0x00, 0x00), (
                        "changed=False 的 run 不应为红色"
                    )

    def test_no_diff_shows_adapted_text(self):
        plan = _make_plan(activity_process_adapted="改写后的过程文本。")
        result = export_daily_plan(plan, [])
        doc = _parse(result)
        cell_text = " ".join(
            p.text for p in doc.tables[0].rows[11].cells[1].paragraphs
        )
        assert "改写后的过程文本" in cell_text

    def test_adapted_newlines_become_separate_paragraphs(self):
        adapted = "一、导入情境，激发兴趣。\n     1. 教师出示图片。\n     2. 引导幼儿观察。"
        plan = _make_plan(activity_process_adapted=adapted)
        result = export_daily_plan(plan, [])
        doc = _parse(result)
        cell = doc.tables[0].rows[11].cells[1]
        # "活动过程："标签独占 1 段，内容 3 行各 1 段 = 共 4 段
        assert len(cell.paragraphs) == 4
        texts = [p.text for p in cell.paragraphs]
        assert any("导入情境" in t for t in texts)
        assert any("教师出示图片" in t for t in texts)
        assert any("引导幼儿观察" in t for t in texts)

    def test_diff_newlines_become_separate_paragraphs_with_color(self):
        diff_result = [
            {"text": "一、导入情境。\n     1. 教师示范。", "changed": True},
            {"text": "二、幼儿操作。", "changed": False},
        ]
        plan = _make_plan()
        result = export_daily_plan(plan, diff_result)
        doc = _parse(result)
        cell = doc.tables[0].rows[11].cells[1]
        all_runs = [run for para in cell.paragraphs for run in para.runs]
        # 验证含"导入情境"的 run 是红色
        red_runs = [r for r in all_runs if r.font.color.rgb == RGBColor(0xFF, 0x00, 0x00)]
        assert any("导入情境" in r.text for r in red_runs)
        # 验证拆分出的第二行("1. 教师示范。")同样标红
        assert any("教师示范" in r.text for r in red_runs)
        # 验证"幼儿操作"不是红色
        normal_runs = [r for r in all_runs if r.font.color.rgb != RGBColor(0xFF, 0x00, 0x00)]
        assert any("幼儿操作" in r.text for r in normal_runs)

    def test_morning_talk_written(self):
        plan = _make_plan(
            morning_talk_topic="谈话主题：今天的天气\n问题设计：\n1.今天是晴天还是阴天？",
            morning_talk_questions=None,
        )
        result = export_daily_plan(plan, [])
        doc = _parse(result)
        topic_text = " ".join(
            p.text for p in doc.tables[0].rows[4].cells[1].paragraphs
        )
        question_text = " ".join(
            p.text for p in doc.tables[0].rows[5].cells[1].paragraphs
        )
        assert "今天的天气" in topic_text
        assert "晴天" in question_text

    def test_morning_activity_spread_across_cells(self):
        plan = _make_plan(
            morning_activity=(
                "体能大循环：\n集体游戏：钻山洞\n自主游戏：搭积木\n"
                "重点指导：钻山洞/搭积木\n活动目标：\n1.锻炼协调\n"
                "指导要点：\n1.注意安全"
            ),
        )
        result = export_daily_plan(plan, [])
        doc = _parse(result)
        cell2 = " ".join(p.text for p in doc.tables[0].rows[2].cells[1].paragraphs)
        cell3 = " ".join(p.text for p in doc.tables[0].rows[3].cells[1].paragraphs)
        assert "钻山洞" in cell2
        assert "搭积木" in cell2
        assert "重点指导" in cell3
        assert "锻炼协调" in cell3
        assert "注意安全" in cell3

    def test_empty_optional_fields_produce_valid_doc(self):
        plan = _make_plan(
            activity_goal=None,
            activity_prep=None,
            activity_key=None,
            activity_difficult=None,
            activity_process_original=None,
            activity_process_adapted=None,
            morning_activity=None,
            indoor_area=None,
            outdoor_activity=None,
            morning_talk_topic=None,
            morning_talk_questions=None,
        )
        result = export_daily_plan(plan, [])
        doc = _parse(result)
        assert len(doc.tables[0].rows) == 19


class TestParseFields:
    def test_parses_labels_and_numbered_items(self):
        text = (
            "游戏区域：美工区 、 建构区\n重点指导：美工区/建构区\n"
            "活动目标：\n1.发展精细动作\n2.激发创造力"
        )
        parsed = _parse_fields(text)
        assert parsed["游戏区域"] == "美工区 、 建构区"
        assert parsed["重点指导"] == "美工区/建构区"
        assert "1.发展精细动作" in parsed["活动目标"]
        assert "2.激发创造力" in parsed["活动目标"]

    def test_empty_text_returns_empty(self):
        assert _parse_fields(None) == {}
        assert _parse_fields("") == {}

    def test_unlabeled_text_ignored(self):
        assert _parse_fields("这是一段没有标签的纯文本") == {}


class TestFallbackFromScratch:
    def test_fallback_has_8_rows(self):
        plan = _make_plan()
        result = _export_from_scratch(plan, [])
        doc = _parse(result)
        assert len(doc.tables[0].rows) == 8

    def test_export_falls_back_when_template_missing(self, monkeypatch, tmp_path):
        # 指向不存在的模板路径，验证降级路径
        monkeypatch.setattr(
            exporter_mod, "TEMPLATE_PATH", tmp_path / "missing.docx"
        )
        plan = _make_plan()
        result = export_daily_plan(plan, [])
        doc = _parse(result)
        assert len(doc.tables[0].rows) == 8


class TestBatchExport:
    def test_batch_empty_list_returns_valid_bytes(self):
        result = export_batch_daily_plans([])
        assert isinstance(result, bytes)
        assert len(result) > 0
        # 确保是合法 docx（能被解析）
        doc = _parse(result)
        assert doc is not None

    def test_batch_single_plan_has_one_table(self):
        plan = _make_plan(plan_date=date(2026, 5, 12))
        result = export_batch_daily_plans([(plan, [])])
        doc = _parse(result)
        assert len(doc.tables) == 1

    def test_batch_multiple_plans_sorted_by_date(self):
        # 故意传入逆序：plan2 日期更早，应排在前面
        plan1 = _make_plan(
            plan_date=date(2026, 5, 20),
            week_number=4,
            weekday_cn="周三",
        )
        plan2 = _make_plan(
            plan_date=date(2026, 5, 12),
            week_number=3,
            weekday_cn="周二",
        )
        result = export_batch_daily_plans([(plan1, []), (plan2, [])])
        doc = _parse(result)
        # 两个 plan → 两张表格
        assert len(doc.tables) == 2
        # 第一张表格应为较早的 plan2（5月12日）
        first_table_text = doc.tables[0].rows[1].cells[0].text
        assert "12" in first_table_text
        # 第二张表格应为较晚的 plan1（5月20日）
        second_table_text = doc.tables[1].rows[1].cells[0].text
        assert "20" in second_table_text
